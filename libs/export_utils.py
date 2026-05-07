# -*- coding: utf-8 -*-
import calendar
import os
import platform
import subprocess

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.worksheet.page import PageMargins

from libs import (
    case_utils,
    date_utils,
    nhi_utils,
    number_utils,
    personnel_utils,
    string_utils,
    system_utils,
)

light_blue = PatternFill(start_color="99CCFF", end_color="99CCFF", fill_type="solid")
light_yellow = PatternFill(start_color="FFFF99", end_color="FFFF99", fill_type="solid")
light_gray = PatternFill(start_color="C0C0C0", end_color="C0C0C0", fill_type="solid")
alternate_row_color = PatternFill(
    start_color="FFC7CE", end_color="FFC7CE", fill_type="solid"
)

salmon = PatternFill(start_color="FFB266", end_color="FFB266", fill_type="solid")
purple = PatternFill(start_color="FF9999", end_color="FF9999", fill_type="solid")
align_center = Alignment(horizontal="center", vertical="center")
align_left = Alignment(horizontal="left", vertical="center")
align_right = Alignment(horizontal="right", vertical="center")
bold = Font(bold=True)
side = Side(border_style="thin", color="000000")
border = Border(top=side, bottom=side, left=side, right=side)


# ====== 加框線（整張表每個 cell）======
def add_border_to_all_cells(sheet):
    thin = Side(border_style="thin", color="000000")  # 黑色細線
    border = Border(top=thin, left=thin, right=thin, bottom=thin)

    for row in sheet.iter_rows(
        min_row=1, max_row=sheet.max_row, min_col=1, max_col=sheet.max_column
    ):
        for cell in row:
            cell.border = border


def export_table_widget_to_excel(
    excel_file_name,
    in_table_widget,
    hidden_column=None,
    numeric_cell=None,
    title=None,
    column_width=None,
    calc_total=False,
    mark_col_no=None,
    variant_col=None,
):
    if numeric_cell is None:
        numeric_cell = []
    wb = Workbook()
    ws = wb.active
    ws.title = "sheet1"

    if column_width is not None:
        for i in range(len(column_width)):
            ws.column_dimensions[chr(ord("A") + i)].width = column_width[i]

    header_row = []
    for col_no in range(in_table_widget.columnCount()):
        if hidden_column is not None and col_no in hidden_column:
            continue

        if in_table_widget.horizontalHeaderItem(col_no) is None:
            continue

        header_row.append(in_table_widget.horizontalHeaderItem(col_no).text())

    if title is not None:
        ws.append([title])

    ws.append(header_row)

    cell_list = string_utils.get_cell_name(end="ZZ")

    start_no = 3
    last_row_no = start_no  # 從row3開始填入資料
    for row_no in range(in_table_widget.rowCount()):
        if in_table_widget.isRowHidden(row_no):
            continue

        if mark_col_no is not None:
            check_box = in_table_widget.cellWidget(row_no, mark_col_no)
            if not check_box.isChecked():
                continue

        last_row_no += 1
        row = []
        for col_no in range(in_table_widget.columnCount()):
            if hidden_column is not None and col_no in hidden_column:
                continue

            item = in_table_widget.item(row_no, col_no)
            item_text = ""
            if item is not None:
                try:
                    item_text = item.text()
                    item_text = string_utils.remove_illegal_characters(item_text)
                except Exception:
                    pass

            if numeric_cell is not None and col_no in numeric_cell:
                item_text = item_text.replace(",", "")
                item_text = number_utils.get_float(item_text)

            if (
                variant_col is not None
                and col_no == variant_col
                and item_text in ["", None]
            ):
                item = in_table_widget.cellWidget(row_no, col_no)
                try:
                    item_text = item.text()
                    item_text = item_text.replace('<br><font size="2" color="red">', "")
                    item_text = item_text.replace("</font>", "")
                except Exception:
                    item_text = ""

            row.append(item_text)

        ws.append(row)

    if calc_total:
        row = []
        for col_no in range(in_table_widget.columnCount()):
            if hidden_column is not None and col_no in hidden_column:
                continue

            item_text = ""
            if numeric_cell is not None and col_no in numeric_cell:
                cell_name = cell_list[col_no - 4]
                item_text = f"=SUM({cell_name}{start_no}:{cell_name}{last_row_no - 1})"

            row.append(item_text)

        ws.append(row)

    wb.save(excel_file_name)

    try:
        # subprocess.Popen([(excel_file_name)], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


def open_file(path):
    if platform.system() == "Windows":
        os.startfile(path)
    elif platform.system() == "Darwin":  # macOS
        subprocess.Popen(["open", path])
    else:  # Linux
        pass
        # subprocess.Popen(["xdg-open", path])


def export_tab_widget_to_excel(
    excel_file_name,
    in_tab_widget,
    hidden_column=None,
    numeric_cell=None,
    title=None,
    column_width=None,
):
    if numeric_cell is None:
        numeric_cell = []

    wb = Workbook()
    ws = wb.active
    for i in range(in_tab_widget.count()):
        current_tab = in_tab_widget.widget(i)
        table_widget = current_tab.tableWidget_medicine_sales

        sheet_name = in_tab_widget.tabText(i)
        if sheet_name in [None, ""]:
            continue

        sheet_name = sheet_name.replace("/", "|")
        ws = wb.create_sheet(sheet_name, i)
        ws.title = sheet_name

        if column_width is not None:
            for i in range(len(column_width)):
                ws.column_dimensions[chr(ord("A") + i)].width = column_width[i]

        header_row = []
        for col_no in range(table_widget.columnCount()):
            if hidden_column is not None and col_no in hidden_column:
                continue

            if table_widget.horizontalHeaderItem(col_no) is None:
                continue

            header_row.append(table_widget.horizontalHeaderItem(col_no).text())

        if title is not None:
            ws.append([title])

        ws.append(header_row)

        for row_no in range(table_widget.rowCount()):
            row = []
            for col_no in range(table_widget.columnCount()):
                if hidden_column is not None and col_no in hidden_column:
                    continue

                item = table_widget.item(row_no, col_no)
                if item is not None:
                    item_text = item.text()
                else:
                    item_text = ""

                if numeric_cell is not None and col_no in numeric_cell:
                    item_text = item_text.replace(",", "")
                    item_text = number_utils.get_float(item_text)

                row.append(item_text)

            ws.append(row)

    wb.save(excel_file_name)

    try:
        # subprocess.Popen([(excel_file_name)], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


def export_multiple_table_widgets_to_excel(
    excel_file_name,
    table_widget_master,
    table_widget_detail,
    hidden_column=None,
    numeric_cell=None,
    title=None,
    column_width=None,
):
    if numeric_cell is None:
        numeric_cell = []

    wb = Workbook()
    ws = wb.active
    for row_no in range(table_widget_master.rowCount()):
        table_widget_master.setCurrentCell(row_no, 0)

        item = table_widget_master.item(row_no, 2)
        if item is None:
            continue

        medicine_type = item.text()
        ws = wb.create_sheet(medicine_type, row_no)
        ws.title = medicine_type

        if column_width is not None:
            for i in range(len(column_width)):
                ws.column_dimensions[chr(ord("A") + i)].width = column_width[i]

        header_row = []
        for col_no in range(table_widget_detail.columnCount()):
            if hidden_column is not None and col_no in hidden_column:
                continue

            if table_widget_detail.horizontalHeaderItem(col_no) is None:
                continue

            header_row.append(table_widget_detail.horizontalHeaderItem(col_no).text())

        if title is not None:
            ws.append([title])

        ws.append(header_row)

        for row_no in range(table_widget_detail.rowCount()):
            row = []
            for col_no in range(table_widget_detail.columnCount()):
                if hidden_column is not None and col_no in hidden_column:
                    continue

                item = table_widget_detail.item(row_no, col_no)
                if item is not None:
                    item_text = item.text()
                else:
                    item_text = ""

                if numeric_cell is not None and col_no in numeric_cell:
                    item_text = item_text.replace(",", "")
                    item_text = number_utils.get_float(item_text)

                row.append(item_text)

            ws.append(row)

    wb.save(excel_file_name)

    try:
        # subprocess.Popen([(excel_file_name)], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


def _get_case_row(row):
    card = row["Card"]
    if card in ["免卡", None]:
        card = ""

    sale_item = ""
    row = [
        row["CaseDate"],
        row["Period"],
        row["Doctor"],
        row["InsType"],
        row["PatientKey"],
        row["Name"],
        row["RegistNo"],
        card,
        row["Course"],
        row["TreatType"],
        row["RegistFee"],
        row["DiagShareFee"],
        row["DrugShareFee"],
        sale_item,
        row["TotalFee"],
        None,
        None,
        0,  # massage_fee
        row["Massager"],
        row["DepositFee"],
        row["DiscountType"],
    ]

    return row


def _get_discount_row(row):
    card = row["Card"]
    if card in ["免卡", None]:
        card = ""

    if row["Massager"] != "":
        sale_item = ""
        massage_item = "折扣"
        massage_fee = -row["DiscountFee"]
        total_fee = 0
    else:
        sale_item = "折扣"
        massage_item = ""
        massage_fee = 0
        total_fee = -row["DiscountFee"]

    row = [
        row["CaseDate"],
        row["Period"],
        row["Doctor"],
        row["InsType"],
        row["PatientKey"],
        row["Name"],
        row["RegistNo"],
        card,
        row["Course"],
        row["RegistFee"],
        row["DiagShareFee"],
        row["DrugShareFee"],
        sale_item,
        total_fee,
        None,
        massage_item,
        massage_fee,
        row["Massager"],
        row["DepositFee"],
        row["DiscountType"],
    ]

    return row, total_fee


# 匯出日報表 From medical_record_list 2019.07.01 板橋新生堂
def export_daily_medical_records_to_excel(
    database, system_settings, excel_file_name, medical_record_rows, show_summary=True
):
    row_range = [
        "A",
        "B",
        "C",
        "D",
        "E",
        "F",
        "G",
        "H",
        "I",
        "J",
        "K",
        "L",
        "M",
        "N",
        "O",
        "P",
        "Q",
        "R",
        "S",
        "T",
        "U",
        "V",
    ]

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "門診日報表資料"

    clinic_name = system_settings.field("院所名稱")
    sheet.append([f"{clinic_name} 日報表"])
    sheet.merge_cells("A1:V1")
    row_property = sheet.row_dimensions[1]
    row_property.height = 30
    row_property.alignment = align_center
    row_property.font = bold

    header_row = [
        "日期",
        "班別",
        "醫師",
        "保險",
        "病歷號",
        "病患姓名",
        "診號",
        "卡序",
        "療程",
        "就醫類別",
        "掛號費",
        "門診負擔",
        "藥品負擔",
        "自費品項",
        "自費金額",
        "推拿金額",
        "推拿自費品項",
        "推拿自費金額",
        "推拿師",
        "押單",
        "備註",
        "合計",
    ]
    sheet.append(header_row)
    sheet.row_dimensions[2].height = 30

    sheet.column_dimensions["A"].width = 14
    sheet.column_dimensions["B"].width = 5
    sheet.column_dimensions["C"].width = 10
    sheet.column_dimensions["D"].width = 5
    sheet.column_dimensions["F"].width = 10
    sheet.column_dimensions["G"].width = 6
    sheet.column_dimensions["H"].width = 6
    sheet.column_dimensions["I"].width = 5

    sheet.column_dimensions["J"].width = 30  # 就醫類別

    sheet.column_dimensions["K"].width = 10
    sheet.column_dimensions["L"].width = 10
    sheet.column_dimensions["M"].width = 10
    sheet.column_dimensions["N"].width = 30
    sheet.column_dimensions["O"].width = 10
    sheet.column_dimensions["P"].width = 10
    sheet.column_dimensions["Q"].width = 30
    sheet.column_dimensions["R"].width = 15
    sheet.column_dimensions["S"].width = 10
    sheet.column_dimensions["V"].width = 13

    total_regist_fee = 0
    total_diag_share_fee = 0
    total_drug_share_fee = 0
    total_purchase_item_fee = 0
    total_massage_item_fee = 0
    total_deposit_fee = 0
    ins_count = 0

    daily_regist_fee = 0
    daily_diag_share_fee = 0
    daily_drug_share_fee = 0
    daily_purchase_item_fee = 0
    daily_massage_item_fee = 0
    daily_deposit_fee = 0
    daily_total = 0
    daily_ins_count = 0

    for row_no, row in enumerate(medical_record_rows):
        case_key = row["CaseKey"]
        case_date = row["CaseDate"]
        period = row["Period"]
        next_case_date = row["NextCaseDate"]
        next_period = row["NextPeriod"]
        ins_type = row["InsType"]
        if ins_type == "健保":
            ins_count += 1

        # treat_type = row['TreatType']
        massager = row["Massager"]
        total_regist_fee += row["RegistFee"]
        total_diag_share_fee += row["DiagShareFee"]
        total_drug_share_fee += row["DrugShareFee"]
        total_deposit_fee += row["DepositFee"]

        if row["DiscountFee"] > 0:
            discount_row, discount_fee = _get_discount_row(row)
            sheet.append(discount_row)
            total_purchase_item_fee += discount_fee

        excel_row = _get_case_row(row)

        sql = f"""
            SELECT * FROM prescript
            WHERE
                CaseKey = {case_key} AND
                MedicineSet >= 2
            ORDER BY PrescriptKey
        """
        prescript_rows = database.select_record(sql)

        if len(prescript_rows) <= 0:
            sheet.append(excel_row)
        else:
            for prescript_row_no, prescript_row in enumerate(prescript_rows):
                if massager == "":
                    pres_days = case_utils.get_pres_days(
                        database, case_key, prescript_row["MedicineSet"]
                    )
                    if pres_days <= 0:
                        pres_days = 1

                    purchase_item = string_utils.xstr(prescript_row["MedicineName"])
                    purchase_item_fee = (
                        number_utils.get_integer(prescript_row["Amount"]) * pres_days
                    )
                    massage_item = ""
                    massage_item_fee = 0
                    total_purchase_item_fee += purchase_item_fee
                else:
                    purchase_item = ""
                    purchase_item_fee = 0
                    massage_item = string_utils.xstr(prescript_row["MedicineName"])
                    massage_item_fee = number_utils.get_integer(prescript_row["Amount"])
                    total_massage_item_fee += massage_item_fee

                excel_row[13] = purchase_item
                excel_row[14] = purchase_item_fee
                excel_row[16] = massage_item
                excel_row[17] = massage_item_fee

                if prescript_row_no > 0:
                    excel_row[3] = ""
                    excel_row[7] = ""
                    excel_row[8] = ""
                    excel_row[9] = ""
                    excel_row[10] = ""
                    excel_row[11] = ""
                    excel_row[12] = ""

                sheet.append(excel_row)

        name = ""
        show_ins_count = ins_count
        if not show_summary:
            show_ins_count = ""
            name = "合計"

        if (case_date == next_case_date and period != next_period) or (
            case_date != next_case_date
        ):
            subtotal = (
                total_regist_fee
                + total_diag_share_fee
                + total_drug_share_fee
                + total_purchase_item_fee
                + total_massage_item_fee
                + total_deposit_fee
            )
            subtotal_row = [
                excel_row[0],
                excel_row[1],
                None,
                show_ins_count,
                None,
                name,
                None,
                None,
                None,
                None,
                total_regist_fee,
                total_diag_share_fee,
                total_drug_share_fee,
                None,
                total_purchase_item_fee,
                None,
                None,
                total_massage_item_fee,
                None,
                total_deposit_fee,
                None,
                subtotal,
            ]
            sheet.append(subtotal_row)

            row_no = sheet._current_row
            for col in row_range:
                cell = sheet[f"{col}{row_no}"]
                cell.fill = salmon
                cell.font = bold

            daily_ins_count += ins_count
            daily_regist_fee += total_regist_fee
            daily_diag_share_fee += total_diag_share_fee
            daily_drug_share_fee += total_drug_share_fee
            daily_purchase_item_fee += total_purchase_item_fee
            daily_massage_item_fee += total_massage_item_fee
            daily_deposit_fee += total_deposit_fee
            daily_total += subtotal

            ins_count = 0
            total_regist_fee = 0
            total_diag_share_fee = 0
            total_drug_share_fee = 0
            total_purchase_item_fee = 0
            total_massage_item_fee = 0
            total_deposit_fee = 0

        if show_summary and case_date != next_case_date:
            sheet.append(
                [
                    excel_row[0],
                    "全",
                    None,
                    daily_ins_count,
                    None,
                    None,
                    None,
                    None,
                    None,
                    None,
                    daily_regist_fee,
                    daily_diag_share_fee,
                    daily_drug_share_fee,
                    None,
                    daily_purchase_item_fee,
                    None,
                    None,
                    daily_massage_item_fee,
                    None,
                    daily_deposit_fee,
                    None,
                    daily_total,
                ]
            )
            row_no = sheet._current_row
            for col in row_range:
                cell = sheet[f"{col}{row_no}"]
                cell.fill = purple
                cell.font = bold

            daily_ins_count = 0
            daily_regist_fee = 0
            daily_diag_share_fee = 0
            daily_drug_share_fee = 0
            daily_purchase_item_fee = 0
            daily_massage_item_fee = 0
            daily_deposit_fee = 0
            daily_total = 0

    if not show_summary:  # 總計 收費證明專用
        total = [
            None,
            None,
            None,
            None,
            None,
            "總計",
            None,
            None,
            None,
            0,
            0,
            0,
            None,
            None,
            daily_purchase_item_fee,
            None,
            None,
            daily_massage_item_fee,
            None,
            0,
            None,
            daily_total,
        ]

        sheet.append(total)
        row_no = sheet._current_row
        for col in row_range:
            cell = sheet[f"{col}{row_no}"]
            cell.fill = purple
            cell.font = bold

    for row_no in ["E2", "F2", "G2", "H2", "I2"]:
        sheet[row_no].fill = light_blue
        sheet[row_no].font = bold
        sheet[row_no].alignment = align_center

    for row_no in ["K2", "L2", "M2"]:
        sheet[row_no].fill = light_yellow
        sheet[row_no].font = bold
        sheet[row_no].alignment = align_center

    for row_no in ["J2", "N2", "O2", "P2", "Q2", "R2", "S2"]:
        sheet[row_no].fill = light_gray
        sheet[row_no].font = bold
        sheet[row_no].alignment = align_center

    for row_no in range(2, len(sheet["A"]) + 1):
        for col in row_range:
            cell = sheet[f"{col}{row_no}"]
            cell.border = border
            cell.alignment = align_center
            if row_no == 2:
                cell.font = bold

    for row_no in [2, 3]:
        for col in row_range + ["V"]:
            cell = sheet[f"{col}{row_no}"]
            sheet.freeze_panes = cell

    workbook.save(excel_file_name)
    try:
        # subprocess.Popen([excel_file_name], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


# 匯出院民照護日報表
def export_nursing_home_list_to_excel(
    database, system_settings, excel_file_name, apply_year, apply_month, clinic_id
):
    last_day = calendar.monthrange(apply_year, apply_month)[1]
    start_date = f"{apply_year}-{apply_month:0>2}-01"
    end_date = f"{apply_year}-{apply_month:0>2}-{last_day:0>2}"

    sql = f'''
        SELECT * FROM cases
        WHERE
            DATE(CaseDate) BETWEEN "{start_date}" AND "{end_date}" AND
            RegistType = "照護機構中醫照護" AND
            InsType = "健保" AND
            ApplyType = "申報" AND
            Card != "欠卡"
        GROUP BY DATE(CaseDate)
    '''
    rows = database.select_record(sql)

    case_count = len(rows)
    if case_count <= 0:
        return

    workbook = Workbook()
    sheet = workbook.active
    workbook.remove(sheet)
    for row in rows:
        add_nursing_home_sheet(
            database, system_settings, row, workbook, apply_year, apply_month, clinic_id
        )

    workbook.save(excel_file_name)
    try:
        # subprocess.Popen([excel_file_name], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


def add_nursing_home_sheet(
    database, system_settings, row, workbook, apply_year, apply_month, clinic_id
):
    apply_date = nhi_utils.get_apply_date(apply_year, apply_month)
    case_date = date_utils.west_date_to_nhi_date(row["CaseDate"], "-")
    title = "全民健康保險中醫門診總額照護機構中醫醫療照護方案門診日報表"
    sheet = workbook.create_sheet(case_date)

    branch_name = system_settings.field("健保業務").split("業務組")[0]
    branch = f"中保會{branch_name}分會"
    case_key = row["CaseKey"]
    sql = f"""
        SELECT TourArea FROM cases
        WHERE
            CaseKey = {case_key}
    """
    rows = database.select_record(sql)

    if len(rows) > 0:
        tour_area = string_utils.xstr(rows[0]["TourArea"])
    else:
        tour_area = ""

    sheet.append(
        [
            title,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            "所　屬　分　會",
            None,
            branch,
        ]
    )
    sheet.append(
        [
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            "承　辦　單　位",
            None,
        ]
    )
    sheet.append(
        [
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            "醫事服務機構代碼",
            None,
            clinic_id,
        ]
    )
    sheet.append(
        [
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            "地　　　　　點",
            None,
            tour_area,
        ]
    )
    sheet.append(
        [
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            "核　准　代　碼",
            None,
        ]
    )

    merge_cell = [
        "A1:S5",
        "T1:U1",
        "T2:U2",
        "T3:U3",
        "T4:U4",
        "T5:U5",
        "V1:X1",
        "V2:X2",
        "V3:X3",
        "V4:X4",
        "V5:X5",
    ]
    for cell in merge_cell:
        sheet.merge_cells(cell)

    for i in range(1, 6):
        row_property = sheet.row_dimensions[i]
        row_property.alignment = align_center
        row_property.font = bold

    title = sheet.cell(row=1, column=1)
    title.font = Font(size=24)
    title.alignment = align_center

    header_row = [
        "日期",
        case_date,
        None,
        "時間",
        "08:00 - 12:00",
        None,
    ]
    sheet.append(header_row)
    sheet.merge_cells("B6:C6")
    sheet.merge_cells("E6:F6")
    sheet.merge_cells("G6:X6")
    row_property = sheet.row_dimensions[6]
    row_property.font = bold
    row_property.alignment = align_center

    header_row = [
        "編號",
        "姓名",
        "身份證統一編號",
        "出生年月日",
        "性別",
        "住址",
        "電話",
        "診察費",
        "藥費(天)",
        "調劑費",
        None,
        "治療處置",
        None,
        None,
        None,
        None,
        None,
        "當地居民",
        None,
        "醫療費用",
        "部份負擔",
        "申請費用",
        "身份別",
        "備註",
    ]
    sheet.append(header_row)
    row_property.font = bold
    sheet.merge_cells("J7:K7")
    sheet.merge_cells("L7:Q7")
    sheet.merge_cells("R7:S7")
    row_property.alignment = align_center

    header_row = [
        None,
        None,
        None,
        None,
        None,
        None,
        None,
        None,
        None,
        "A31",
        "A32",
        "1",
        "2",
        "3",
        "4",
        "5",
        "6",
        "是",
        "否",
    ]
    sheet.append(header_row)
    merge_cell = [
        "A7:A8",
        "B7:B8",
        "C7:C8",
        "D7:D8",
        "E7:E8",
        "F7:F8",
        "G7:G8",
        "H7:H8",
        "I7:I8",
        "T7:T8",
        "U7:U8",
        "V7:V8",
        "W7:W8",
        "X7:X8",
    ]
    for cell in merge_cell:
        sheet.merge_cells(cell)

    row_property = sheet.row_dimensions[7]
    row_property.alignment = align_center
    row_property.font = bold
    row_property = sheet.row_dimensions[8]
    row_property.alignment = align_center
    row_property.font = bold

    sheet.column_dimensions["B"].width = 10
    sheet.column_dimensions["C"].width = 13
    sheet.column_dimensions["D"].width = 13
    sheet.column_dimensions["F"].width = 35
    sheet.column_dimensions["G"].width = 15
    adjust_column = [
        "E",
        "J",
        "K",
        "L",
        "M",
        "N",
        "O",
        "P",
        "Q",
        "R",
        "S",
    ]
    for cell in adjust_column:
        sheet.column_dimensions[cell].width = 5

    case_date = row["CaseDate"]
    sql = f'''
        SELECT insapply.*,
               cases.TourArea,
               patient.Gender, patient.Address, patient.Telephone
        FROM insapply
            LEFT JOIN cases ON insapply.CaseKey1 = cases.CaseKey
            LEFT JOIN patient ON insapply.PatientKey = patient.PatientKey
        WHERE
            ApplyDate = "{apply_date}" AND
            insapply.ApplyType = 1 AND
            ApplyPeriod = "全月" AND
            CaseType = "22" AND
            cases.RegistType = "照護機構中醫照護" AND
            insapply.CaseDate = "{case_date.date()}"
        ORDER BY Sequence
    '''
    rows = database.select_record(sql)

    treatment_stats = {}
    start_row = 9  # 資料開始的行數

    for row_no, row in enumerate(rows):
        if string_utils.xstr(row["Gender"]) == "男":
            gender_code = "1"
        elif string_utils.xstr(row["Gender"]) == "女":
            gender_code = "0"
        else:
            gender_code = ""

        pres_days = row["PresDays"]
        pharmacy_code = nhi_utils.extract_pharmacy_code(
            string_utils.xstr(row["PharmacyCode"])
        )
        pharmacy_list = [None, None]
        if pharmacy_code != "":
            pharmacy_code_dict = {
                "A31": 0,
                "A32": 1,
            }
            pharmacy_list[pharmacy_code_dict[pharmacy_code]] = "V"

        treat_list = [None] * 6
        treat_code = string_utils.xstr(row["TreatCode1"])
        if treat_code != "":
            treat_list[0] = treat_code
            treatment_stats[treat_code] = treatment_stats.get(treat_code, 0) + 1

        address = string_utils.xstr(row["Address"])
        native_list = [
            None,
            None,
        ]
        if address == "" or tour_area in address:
            native_list[0] = "V"
        else:
            native_list[1] = "V"

        share_code = string_utils.xstr(row["ShareCode"])
        if share_code in ["S10", "S20"]:
            share_code = ""

        def to_num(val):
            try:
                return float(val) if val is not None else 0
            except (ValueError, TypeError):
                return 0

        data = [
            row_no + 1,
            string_utils.xstr(row["Name"]),
            string_utils.xstr(row["ID"]),
            string_utils.xstr(date_utils.west_date_to_nhi_date(row["Birthday"], "-")),
            gender_code,
            address,
            string_utils.xstr(row["Telephone"]),
            string_utils.xstr(row["DiagCode"]),
            pres_days,
            pharmacy_list[0],
            pharmacy_list[1],
            treat_list[0],
            treat_list[1],
            treat_list[2],
            treat_list[3],
            treat_list[4],
            treat_list[5],
            native_list[0],
            native_list[1],
            to_num(row["InsTotalFee"]),  # 轉為數字 (T欄)
            to_num(row["ShareFee"]),  # 轉為數字 (U欄)
            to_num(row["InsApplyFee"]),  # 轉為數字 (V欄)
            share_code,
        ]
        sheet.append(data)

        this_row_idx = row_no + start_row

        # 針對該列的每一個儲存格進行設定
        for col_idx in range(1, len(data) + 1):
            cell = sheet.cell(row=this_row_idx, column=col_idx)
            if col_idx in [20, 21, 22]:  # T, U, V
                cell.alignment = Alignment(horizontal="right", vertical="center")
            else:
                cell.alignment = align_center

    # --- 關鍵修改：排序與輸出合計 ---
    summary_row_no = len(rows) + start_row  # 合計字樣所在列

    # 寫入「合計」字樣
    sheet.cell(row=summary_row_no, column=11).value = "合計"
    sheet.cell(row=summary_row_no, column=11).font = bold
    sheet.cell(row=summary_row_no, column=11).alignment = align_center

    # 2. 將統計字典依照 Key (代碼字母) 進行排序
    # sorted() 會依據字母 A-Z, 數字 0-9 排序
    sorted_codes = sorted(treatment_stats.items())

    # 寫入合計金額 (T, U, V 欄)
    for col in [20, 21, 22]:
        cell = sheet.cell(row=summary_row_no, column=col)
        # 這裡同樣設定為靠右
        cell.alignment = Alignment(horizontal="right", vertical="center")
        cell.font = bold
        # 填入公式
        col_letter = "T" if col == 20 else "U" if col == 21 else "V"
        cell.value = f"=SUM({col_letter}9:{col_letter}{summary_row_no - 1})"

    for i, (code, count) in enumerate(sorted_codes):
        current_col = 12 + i  # 從 L 欄 (12) 開始向右填寫
        # 上列：代碼
        sheet.cell(row=summary_row_no, column=current_col).value = code
        sheet.cell(row=summary_row_no, column=current_col).font = bold
        # 下列：數量
        sheet.cell(row=summary_row_no + 1, column=current_col).value = count

        # 設定居中
        sheet.cell(row=summary_row_no, column=current_col).alignment = align_center
        sheet.cell(row=summary_row_no + 1, column=current_col).alignment = align_center

    # 金額總計公式 (T, U, V 欄)
    for col in [20, 21, 22]:
        col_letter = "T" if col == 20 else "U" if col == 21 else "V"
        cell = sheet.cell(row=summary_row_no, column=col)
        cell.value = f"=SUM({col_letter}{start_row}:{col_letter}{summary_row_no - 1})"
        cell.alignment = Alignment(horizontal="right", vertical="center")
        cell.font = bold


def get_period(database, case_date):
    current_date = date_utils.str_to_date(case_date)
    weekday = date_utils.WEEK_DAY_LIST[current_date.weekday()]
    start_time = ""
    end_time = ""

    sql = f"""
        SELECT * FROM doctor_schedule
        WHERE
            Period = "早班" AND
            {weekday} IS NOT NULL
    """
    rows = database.select_record(sql)
    if rows:
        start_time = "09:00"
        end_time = "12:00"

    sql = f"""
        SELECT * FROM doctor_schedule
        WHERE
            Period = "午班" AND
            {weekday} IS NOT NULL
    """
    rows = database.select_record(sql)
    if rows:
        if start_time == "":
            start_time = "13:00"

        end_time = "18:00"

    sql = f"""
        SELECT * FROM doctor_schedule
        WHERE
            Period = "晚班" AND
            {weekday} IS NOT NULL
    """
    rows = database.select_record(sql)
    if rows:
        if start_time == "":
            start_time = "18:00"

        end_time = "22:00"

    return f"{start_time} - {end_time}"


# 匯出巡迴醫療日報表
# def export_tour_daily_list_to_excel(
#         database, system_settings, excel_file_name,
#         apply_date, apply_year, apply_month, apply_type_code, period, clinic_id,
#         tour_type, monthly_list=None):
#     sql = f'''
#         SELECT * FROM insapply
#         WHERE
#             ApplyDate = "{apply_date}" AND
#             ApplyType = "{apply_type_code}" AND
#             ApplyPeriod = "{period}" AND
#             ClinicID = "{clinic_id}" AND
#             CaseType = "25"
#         GROUP BY CaseDate
#     '''
#     rows = database.select_record(sql)
#     tour_apply_count = len(rows)
#     if tour_apply_count <= 0:
#         return

#     if monthly_list is not None:
#         rows = monthly_list

#     workbook = Workbook()
#     sheet = workbook.active
#     workbook.remove_sheet(sheet)
#     for row in rows:
#         add_tour_daily_list_sheet(
#             database, system_settings, row, workbook, apply_year, apply_month,
#             apply_type_code, period, clinic_id, tour_type
#         )

#     workbook.save(excel_file_name)
#     try:
#         subprocess.Popen([excel_file_name], shell=True)
#     except Exception:
#         pass


# def add_tour_daily_list_sheet(
#         database, system_settings, row, workbook,
#         apply_year, apply_month, apply_type_code, period, clinic_id, tour_type):
#     apply_date = nhi_utils.get_apply_date(apply_year, apply_month)
#     case_date = date_utils.west_date_to_nhi_date(row['CaseDate'], '-')
#     title = f'{case_date}{tour_type}門診日報表資料'
#     sheet = workbook.create_sheet(title)

#     branch_name = system_settings.field('健保業務').split('業務組')[0]
#     branch = f'中保會{branch_name}分會'
#     case_key = row['CaseKey1']
#     sql = f'''
#         SELECT TourArea FROM cases
#         WHERE
#             CaseKey = {case_key}
#     '''
#     rows = database.select_record(sql)

#     if len(rows) > 0:
#         tour_area = string_utils.xstr(rows[0]['TourArea'])
#     else:
#         tour_area = system_settings.field('院所地址')[:6]

#     sheet.append([
#         f'{apply_year-1911}年度全民健康保險中醫門診總額醫療資源不足地區醫療服務門診日報表',
#         None, None,
#         None, None, None, None, None, None, None, None, None, None,
#         None, None, None, None, None, None,
#         '所　屬　分　會', None, branch,
#     ])
#     sheet.append([
#         None, None,
#         None, None, None, None, None, None, None, None, None, None,
#         None, None, None, None, None, None, None,
#         '承　辦　單　位', None, system_settings.field('院所名稱'),
#     ])
#     sheet.append([
#         None, None,
#         None, None, None, None, None, None, None, None, None, None,
#         None, None, None, None, None, None, None,
#         '醫事服務機構代碼', None, clinic_id,
#     ])
#     sheet.append([
#         None, None,
#         None, None, None, None, None, None, None, None, None, None,
#         None, None, None, None, None, None, None,
#         '地　　　　　點', None, tour_area,
#     ])
#     sheet.append([
#         None, None,
#         None, None, None, None, None, None, None, None, None, None,
#         None, None, None, None, None, None, None,
#         '核　准　代　碼', None,
#     ])
#     merge_cell = [
#         'A1:S5',
#         'T1:U1', 'T2:U2', 'T3:U3', 'T4:U4', 'T5:U5',
#         'V1:X1', 'V2:X2', 'V3:X3', 'V4:X4', 'V5:X5',
#     ]
#     for cell in merge_cell:
#         sheet.merge_cells(cell)

#     for i in range(1, 6):
#         row_property = sheet.row_dimensions[i]
#         row_property.alignment = align_center
#         row_property.font = bold

#     if tour_type == '資源不足開業':
#         duration = get_period(database, row['CaseDate'])
#     else:
#         duration = '08:00 - 18:00'

#     header_row = [
#         '日期', case_date, None, '時間', duration, None,
#     ]
#     sheet.append(header_row)
#     sheet.merge_cells('B6:C6')
#     sheet.merge_cells('E6:F6')
#     row_property = sheet.row_dimensions[6]
#     row_property.font = bold
#     row_property.alignment = align_center

#     header_row = [
#         '編號', '姓名', '身份證統一編號', '出生年月日', '性別', '住址', '電話',
#         '診察費', '藥費(天)', '調劑費', None,
#         '治療處置',
#         None, None, None, None, None,
#         '當地居民', None,
#         '醫療費用', '部份負擔', '申請費用', '身份別', '備註',
#     ]
#     sheet.append(header_row)
#     row_property.font = bold
#     sheet.merge_cells('J7:K7')
#     sheet.merge_cells('L7:Q7')
#     sheet.merge_cells('R7:S7')
#     row_property.alignment = align_center

#     header_row = [
#         None, None, None, None, None, None, None,
#         None, None, 'A31', 'A32',
#         '1', '2', '3', '4', '5', '6',
#         '是', '否',
#     ]
#     sheet.append(header_row)
#     merge_cell = [
#         'A7:A8', 'B7:B8', 'C7:C8', 'D7:D8', 'E7:E8', 'F7:F8', 'G7:G8',
#         'H7:H8', 'I7:I8',
#         'T7:T8', 'U7:U8', 'V7:V8', 'W7:W8', 'X7:X8',
#     ]
#     for cell in merge_cell:
#         sheet.merge_cells(cell)

#     row_property = sheet.row_dimensions[7]
#     row_property.alignment = align_center
#     row_property.font = bold
#     row_property = sheet.row_dimensions[8]
#     row_property.alignment = align_center
#     row_property.font = bold

#     sheet.column_dimensions['B'].width = 10
#     sheet.column_dimensions['C'].width = 13
#     sheet.column_dimensions['D'].width = 13
#     sheet.column_dimensions['F'].width = 35
#     sheet.column_dimensions['G'].width = 15
#     adjust_column = [
#         'E',
#         'J', 'K',
#         'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S',
#     ]
#     for cell in adjust_column:
#         sheet.column_dimensions[cell].width = 5

#     case_date = row['CaseDate']
#     sql = f'''
#         SELECT *,
#                cases.TourArea,
#                patient.Gender, patient.Address, patient.Telephone
#         FROM insapply
#             LEFT JOIN cases ON insapply.CaseKey1 = cases.CaseKey
#             LEFT JOIN patient ON insapply.PatientKey = patient.PatientKey
#         WHERE
#             ApplyDate = "{apply_date}" AND
#             insapply.ApplyType = "{apply_type_code}" AND
#             ApplyPeriod = "{period}" AND
#             ClinicID = "{clinic_id}" AND
#             CaseType = "25" AND
#             insapply.CaseDate = "{case_date}"
#         ORDER BY Sequence
#     '''
#     rows = database.select_record(sql)

#     for row_no, row in enumerate(rows):
#         if string_utils.xstr(row['Gender']) == '男':
#             gender_code = '1'
#         elif string_utils.xstr(row['Gender']) == '女':
#             gender_code = '0'
#         else:
#             gender_code = ''

#         pres_days = row['PresDays']
#         pharmacy_code = nhi_utils.extract_pharmacy_code(string_utils.xstr(row['PharmacyCode']))
#         pharmacy_list = [None, None]
#         if pharmacy_code != '':
#             pharmacy_code_dict = {
#                 'A31': 0, 'A32': 1,
#             }
#             pharmacy_list[pharmacy_code_dict[pharmacy_code]] = 'V'

#         treat_list = [
#             None, None, None, None, None, None,
#         ]
#         treat_code = string_utils.xstr(row['TreatCode1'])
#         if treat_code != '':
#             treat_list[0] = treat_code

#         address = string_utils.xstr(row['Address'])
#         native_list = [
#             None, None,
#         ]
#         if address == '' or tour_area in address:
#             native_list[0] = 'V'
#         else:
#             native_list[1] = 'V'

#         share_code = string_utils.xstr(row['ShareCode'])
#         if share_code in ['S10', 'S20']:
#             share_code = ''

#         data = [
#             row_no+1,
#             string_utils.xstr(row['Name']),
#             string_utils.xstr(row['ID']),
#             string_utils.xstr(date_utils.west_date_to_nhi_date(row['Birthday'], '-')),
#             gender_code,
#             address,
#             string_utils.xstr(row['Telephone']),
#             string_utils.xstr(row['DiagCode']),
#             pres_days, pharmacy_list[0], pharmacy_list[1],
#             treat_list[0], treat_list[1], treat_list[2], treat_list[3], treat_list[4], treat_list[5],
#             native_list[0], native_list[1],
#             string_utils.xstr(row['InsTotalFee']),
#             string_utils.xstr(row['ShareFee']),
#             string_utils.xstr(row['InsApplyFee']),
#             share_code,
#         ]
#         sheet.append(data)
#         row_property = sheet.row_dimensions[row_no+9]
#         row_property.alignment = align_center
# ====== 匯出多工作表：修正 remove_sheet 棄用用法 ======

# ====== imports（若你原本已有，重複匯入無妨）======
# 通用樣式（若你的專案已有 align_center / bold，可刪除這兩行）


# ====== 小工具：設定 A4 橫向 + 縮放單頁 ======
def set_sheet_to_a4_onepage(
    sheet,
    last_row=None,
    last_col=None,
    orientation="landscape",  # 預設改成橫向
    margin_cm=1.0,
):
    """
    將工作表設定為 A4 橫向、縮放到單頁列印，並限制列印範圍。
    margin_cm：四邊頁邊距（公分）
    """
    # A4 + 方向
    sheet.page_setup.paperSize = sheet.PAPERSIZE_A4
    sheet.page_setup.orientation = orientation

    # 縮放至單頁
    sheet.page_setup.fitToPage = True
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 1  # 若允許跨多頁高，改成 0

    # 頁邊距（openpyxl 使用英吋）
    cm_to_in = 0.3937007874
    m = margin_cm * cm_to_in
    sheet.page_margins = PageMargins(
        left=m, right=m, top=m, bottom=m, header=0.3, footer=0.3
    )

    # 列印範圍
    if last_row is None:
        last_row = sheet.max_row
    if last_col is None:
        last_col = sheet.max_column
    sheet.print_area = f"A1:{get_column_letter(last_col)}{last_row}"

    # 如果未來有跨頁，這行可讓表頭每頁重複（目前單頁可不設）
    # sheet.print_title_rows = "1:8"


# ====== 若資料太少，自動補空白列（例如至少保留 16 列含資料）======
def fill_blank_rows(sheet, start_row, total_rows=16):
    """
    若目前資料筆數不足 total_rows，補足空白列到指定數量。
    """
    current_rows = sheet.max_row - start_row + 1
    if current_rows < total_rows:
        blanks_needed = total_rows - current_rows
        for _ in range(blanks_needed):
            sheet.append([None] * sheet.max_column)


def export_tour_daily_list_to_excel(
    database,
    system_settings,
    excel_file_name,
    apply_date,
    apply_year,
    apply_month,
    apply_type_code,
    period,
    clinic_id,
    tour_type,
    monthly_list=None,
):

    sql = f'''
        SELECT * FROM insapply
        WHERE
            ApplyDate = "{apply_date}" AND
            ApplyType = "{apply_type_code}" AND
            ApplyPeriod = "{period}" AND
            ClinicID = "{clinic_id}" AND
            CaseType = "25"
        GROUP BY CaseDate
    '''
    rows = database.select_record(sql)
    tour_apply_count = len(rows)
    if tour_apply_count <= 0:
        return

    if monthly_list is not None:
        rows = monthly_list

    workbook = Workbook()
    # 取代舊版的 remove_sheet：刪除預設空白工作表
    workbook.remove(workbook.active)

    for row in rows:
        add_tour_daily_list_sheet(
            database,
            system_settings,
            row,
            workbook,
            apply_year,
            apply_month,
            apply_type_code,
            period,
            clinic_id,
            tour_type,
        )

    workbook.save(excel_file_name)
    try:
        # subprocess.Popen([excel_file_name], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


# ====== 建表函式：結尾加入「橫向 A4 單頁列印」設定 ======
def add_tour_daily_list_sheet(
    database,
    system_settings,
    row,
    workbook,
    apply_year,
    apply_month,
    apply_type_code,
    period,
    clinic_id,
    tour_type,
):

    apply_date = nhi_utils.get_apply_date(apply_year, apply_month)
    case_date = date_utils.west_date_to_nhi_date(row["CaseDate"], "-")
    title = f"{case_date}{tour_type}門診日報表資料"
    sheet = workbook.create_sheet(title)

    branch_name = system_settings.field("健保業務").split("業務組")[0]
    branch = f"中保會{branch_name}分會"
    case_key = row["CaseKey1"]
    sql = f"""
        SELECT TourArea FROM cases
        WHERE
            CaseKey = {case_key}
    """
    rows = database.select_record(sql)

    if len(rows) > 0:
        tour_area = string_utils.xstr(rows[0]["TourArea"])
    else:
        tour_area = system_settings.field("院所地址")[:6]

    sheet.append(
        [
            f"{apply_year - 1911}年度全民健康保險中醫門診總額醫療資源不足地區醫療服務門診日報表",
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            "所　屬　分　會",
            None,
            branch,
        ]
    )
    sheet.append(
        [
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            "承　辦　單　位",
            None,
            system_settings.field("院所名稱"),
        ]
    )
    sheet.append(
        [
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            "醫事服務機構代碼",
            None,
            clinic_id,
        ]
    )
    sheet.append(
        [
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            "地　　　　　點",
            None,
            tour_area,
        ]
    )
    sheet.append(
        [
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            "核　准　代　碼",
            None,
        ]
    )
    merge_cell = [
        "A1:S5",
        "T1:U1",
        "T2:U2",
        "T3:U3",
        "T4:U4",
        "T5:U5",
        "V1:X1",
        "V2:X2",
        "V3:X3",
        "V4:X4",
        "V5:X5",
    ]
    for cell in merge_cell:
        sheet.merge_cells(cell)

    sheet["A1"].alignment = Alignment(
        horizontal="center", vertical="center", wrap_text=True
    )
    sheet["A1"].font = Font(bold=True, size=14)

    for i in range(1, 6):
        row_property = sheet.row_dimensions[i]
        row_property.alignment = align_center
        row_property.font = bold

    if tour_type == "資源不足開業":
        duration = get_period(database, row["CaseDate"])
    else:
        duration = "08:00 - 18:00"

    header_row = [
        "日期",
        case_date,
        None,
        "時間",
        duration,
        None,
    ]
    sheet.append(header_row)
    sheet.merge_cells("B6:C6")
    sheet.merge_cells("E6:F6")
    row_property = sheet.row_dimensions[6]
    row_property.font = bold
    row_property.alignment = align_center

    header_row = [
        "編號",
        "姓名",
        "身份證統一編號",
        "出生年月日",
        "性別",
        "住址",
        "電話",
        "診察費",
        "藥費(天)",
        "調劑費",
        None,
        "治療處置",
        None,
        None,
        None,
        None,
        None,
        "當地居民",
        None,
        "醫療費用",
        "部份負擔",
        "申請費用",
        "身份別",
        "備註",
    ]
    sheet.append(header_row)
    row_property.font = bold
    sheet.merge_cells("J7:K7")
    sheet.merge_cells("L7:Q7")
    sheet.merge_cells("R7:S7")
    row_property.alignment = align_center

    header_row = [
        None,
        None,
        None,
        None,
        None,
        None,
        None,
        None,
        None,
        "A31",
        "A32",
        "1",
        "2",
        "3",
        "4",
        "5",
        "6",
        "是",
        "否",
    ]
    sheet.append(header_row)
    merge_cell = [
        "A7:A8",
        "B7:B8",
        "C7:C8",
        "D7:D8",
        "E7:E8",
        "F7:F8",
        "G7:G8",
        "H7:H8",
        "I7:I8",
        "T7:T8",
        "U7:U8",
        "V7:V8",
        "W7:W8",
        "X7:X8",
    ]
    for cell in merge_cell:
        sheet.merge_cells(cell)

    row_property = sheet.row_dimensions[7]
    row_property.alignment = align_center
    row_property.font = bold
    row_property = sheet.row_dimensions[8]
    row_property.alignment = align_center
    row_property.font = bold

    sheet.column_dimensions["B"].width = 10
    sheet.column_dimensions["C"].width = 13
    sheet.column_dimensions["D"].width = 13
    sheet.column_dimensions["F"].width = 35
    sheet.column_dimensions["G"].width = 15
    adjust_column = ["E", "J", "K", "L", "M", "N", "O", "P", "Q", "R", "S"]
    for cell in adjust_column:
        sheet.column_dimensions[cell].width = 5

    case_date = row["CaseDate"]
    sql = f'''
        SELECT *,
               cases.TourArea,
               patient.Gender, patient.Address, patient.Telephone
        FROM insapply
            LEFT JOIN cases ON insapply.CaseKey1 = cases.CaseKey
            LEFT JOIN patient ON insapply.PatientKey = patient.PatientKey
        WHERE
            ApplyDate = "{apply_date}" AND
            insapply.ApplyType = "{apply_type_code}" AND
            ApplyPeriod = "{period}" AND
            ClinicID = "{clinic_id}" AND
            CaseType = "25" AND
            insapply.CaseDate = "{case_date}"
        ORDER BY Sequence
    '''
    rows = database.select_record(sql)

    for row_no, row in enumerate(rows):
        if string_utils.xstr(row["Gender"]) == "男":
            gender_code = "1"
        elif string_utils.xstr(row["Gender"]) == "女":
            gender_code = "0"
        else:
            gender_code = ""

        pres_days = row["PresDays"]
        pharmacy_code = nhi_utils.extract_pharmacy_code(
            string_utils.xstr(row["PharmacyCode"])
        )
        pharmacy_list = [None, None]
        if pharmacy_code != "":
            pharmacy_code_dict = {"A31": 0, "A32": 1}
            pharmacy_list[pharmacy_code_dict[pharmacy_code]] = "V"

        treat_list = [None, None, None, None, None, None]
        treat_code = string_utils.xstr(row["TreatCode1"])
        if treat_code != "":
            treat_list[0] = treat_code

        address = string_utils.xstr(row["Address"])
        native_list = [None, None]
        if address == "" or tour_area in address:
            native_list[0] = "V"
        else:
            native_list[1] = "V"

        share_code = string_utils.xstr(row["ShareCode"])
        if share_code in ["S10", "S20"]:
            share_code = ""

        data = [
            row_no + 1,
            string_utils.xstr(row["Name"]),
            string_utils.xstr(row["ID"]),
            string_utils.xstr(date_utils.west_date_to_nhi_date(row["Birthday"], "-")),
            gender_code,
            address,
            string_utils.xstr(row["Telephone"]),
            string_utils.xstr(row["DiagCode"]),
            pres_days,
            pharmacy_list[0],
            pharmacy_list[1],
            treat_list[0],
            treat_list[1],
            treat_list[2],
            treat_list[3],
            treat_list[4],
            treat_list[5],
            native_list[0],
            native_list[1],
            string_utils.xstr(row["InsTotalFee"]),
            string_utils.xstr(row["ShareFee"]),
            string_utils.xstr(row["InsApplyFee"]),
            share_code,
        ]
        sheet.append(data)
        row_property = sheet.row_dimensions[row_no + 9]
        row_property.alignment = align_center

    # ====== 若只有少數幾筆資料，自動補足 15 個空白列 ======
    # 你的表格實際資料是從第 9 列開始（第 8 列以前是表頭）
    fill_blank_rows(sheet, start_row=9, total_rows=9 + 15)

    add_border_to_all_cells(sheet)
    # ====== 關鍵：把這張表設定為「橫向 A4、縮放到單頁」======
    set_sheet_to_a4_onepage(
        sheet,
        last_row=sheet.max_row,
        last_col=sheet.max_column,
        orientation="landscape",  # 橫向
        margin_cm=1.0,  # 需要更窄可調 0.7 或 0.5
    )


#  綜合業績報表 2020.03.17
def export_multiple_performance_to_excel(
    clinic_name, year, month, excel_file_name, tab_widget
):
    workbook = Workbook()
    sheet = workbook.active
    workbook.remove_sheet(sheet)

    _export_week_table(
        workbook,
        "週人數統計表",
        clinic_name,
        year,
        month,
        tab_widget.widget(0).ui.tableWidget_medical_record,
    )
    _export_week_table(
        workbook,
        "自費收入統計表",
        clinic_name,
        year,
        month,
        tab_widget.widget(1).ui.tableWidget_medical_record,
    )
    _export_week_project(
        workbook,
        "週專案統計表",
        clinic_name,
        year,
        month,
        tab_widget.widget(2).ui.tableWidget_medical_record,
    )
    _export_week_daily(
        workbook,
        "週人數日報表",
        clinic_name,
        year,
        month,
        tab_widget.widget(0).ui.tableWidget_week,
    )
    _export_week_daily(
        workbook,
        "週自費收入日報表",
        clinic_name,
        year,
        month,
        tab_widget.widget(1).ui.tableWidget_week,
    )
    _export_week_doctor(
        workbook,
        "週醫師業績日報表",
        clinic_name,
        year,
        month,
        tab_widget.widget(3).ui.tableWidget_medical_record,
    )

    workbook.save(excel_file_name)
    try:
        # subprocess.Popen([excel_file_name], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


def _export_week_table(workbook, title, clinic_name, year, month, table_widget):
    sheet = workbook.create_sheet(title)

    col_range = [
        "A",
        "B",
        "C",
        "D",
        "E",
        "F",
        "G",
        "H",
    ]
    for col_no in col_range:
        sheet.column_dimensions[col_no].width = 16
    sheet.column_dimensions["A"].width = 18

    row_index = 1
    sheet.row_dimensions[row_index].height = 30
    sheet.row_dimensions[row_index].alignment = align_left
    sheet.row_dimensions[row_index].font = bold
    sheet.append([f"{clinic_name} {year}年{month}月份 {title}"])
    sheet.merge_cells(f"A{row_index}:H{row_index}")

    row_index += 1
    header_row = []
    for i in range(table_widget.columnCount()):
        header_row.append(
            ""
            if table_widget.horizontalHeaderItem(i) is None
            else table_widget.horizontalHeaderItem(i).text()
        )
    sheet.append(header_row)

    header_row = [
        None,
        "" if table_widget.item(0, 1) is None else table_widget.item(0, 1).text(),
        "" if table_widget.item(0, 2) is None else table_widget.item(0, 2).text(),
        "" if table_widget.item(0, 3) is None else table_widget.item(0, 3).text(),
        "" if table_widget.item(0, 4) is None else table_widget.item(0, 4).text(),
        "" if table_widget.item(0, 5) is None else table_widget.item(0, 5).text(),
    ]
    sheet.append(header_row)
    sheet.merge_cells(f"A{row_index}:A{row_index + 1}")
    sheet.merge_cells(f"G{row_index}:G{row_index + 1}")
    sheet.merge_cells(f"H{row_index}:H{row_index + 1}")

    color_header = [
        f"A{row_index}",
        f"B{row_index}",
        f"C{row_index}",
        f"D{row_index}",
        f"E{row_index}",
        f"F{row_index}",
        f"G{row_index}",
        f"H{row_index}",
        f"A{row_index + 1}",
        f"B{row_index + 1}",
        f"C{row_index + 1}",
        f"D{row_index + 1}",
        f"E{row_index + 1}",
        f"F{row_index + 1}",
        f"G{row_index + 1}",
        f"H{row_index + 1}",
    ]

    row_index += 1
    for col_no in color_header:
        sheet[col_no].fill = light_gray
        sheet[col_no].border = border
        sheet[col_no].alignment = align_center
        sheet[col_no].font = bold

    for row_no in range(1, table_widget.rowCount()):
        row = []
        for col_no in range(table_widget.columnCount()):
            row.append(
                ""
                if table_widget.item(row_no, col_no) is None
                else table_widget.item(row_no, col_no).text(),
            )
        sheet.append(row)

        row_index += 1
        for col in col_range:
            cell = sheet[f"{col}{row_index}"]
            cell.border = border
            cell.alignment = align_center
            if row_index % 2 == 0:
                cell.fill = alternate_row_color


def _export_week_project(workbook, title, clinic_name, year, month, table_widget):
    sheet = workbook.create_sheet(title)

    basic_col_range = [
        "A",
        "B",
        "C",
        "D",
        "E",
        "F",
        "G",
        "H",
        "I",
        "J",
        "K",
        "L",
        "M",
        "N",
        "O",
        "P",
        "Q",
        "R",
        "S",
        "T",
        "U",
        "V",
        "W",
        "X",
        "Y",
        "Z",
    ]
    for col_no in basic_col_range:
        sheet.column_dimensions[col_no].width = 15
    sheet.column_dimensions["A"].width = 10

    row_index = 1
    sheet.row_dimensions[row_index].height = 30
    sheet.row_dimensions[row_index].alignment = align_left
    sheet.row_dimensions[row_index].font = bold
    sheet.append([f"{clinic_name} {year}年{month}月份 {title}"])
    sheet.merge_cells(f"A{row_index}:H{row_index}")

    row_index += 1
    col_range = []
    color_header = []
    for col_no in range(table_widget.columnCount()):
        col_range.append(basic_col_range[col_no])
        color_header.append(f"{basic_col_range[col_no]}{row_index}")

    header_row = []
    for i in range(table_widget.columnCount()):
        header_row.append(
            ""
            if table_widget.horizontalHeaderItem(i) is None
            else table_widget.horizontalHeaderItem(i).text()
        )
    sheet.append(header_row)

    for col_no in color_header:
        sheet[col_no].fill = light_gray
        sheet[col_no].border = border
        sheet[col_no].alignment = align_center
        sheet[col_no].font = bold

    for row_no in range(table_widget.rowCount()):
        row = []
        for col_no in range(table_widget.columnCount()):
            row.append(
                ""
                if table_widget.item(row_no, col_no) is None
                else table_widget.item(row_no, col_no).text(),
            )
        sheet.append(row)

        row_index += 1
        for col in col_range:
            cell = sheet[f"{col}{row_index}"]
            cell.border = border
            cell.alignment = align_center
            if row_index % 2 == 0:
                cell.fill = alternate_row_color


def _export_week_daily(workbook, title, clinic_name, year, month, table_widget):
    sheet = workbook.create_sheet(title)

    col_range = [
        "A",
        "B",
        "C",
        "D",
        "E",
        "F",
        "G",
        "H",
        "I",
        "J",
        "K",
        "L",
        "M",
        "N",
        "O",
        "P",
        "Q",
        "R",
        "S",
        "T",
        "U",
        "V",
        "W",
    ]
    for col_no in col_range:
        sheet.column_dimensions[col_no].width = 10

    sheet.column_dimensions["A"].width = 8

    row_index = 1
    sheet.row_dimensions[row_index].height = 30
    sheet.row_dimensions[row_index].alignment = align_left
    sheet.row_dimensions[row_index].font = bold
    sheet.append([f"{clinic_name} {year}年{month}月份 {title}"])
    sheet.merge_cells(f"A{row_index}:W{row_index}")

    row_index += 1
    header_row = [
        "週數",
        "星期一",
        None,
        None,
        "星期二",
        None,
        None,
        "星期三",
        None,
        None,
        "星期四",
        None,
        None,
        "星期五",
        None,
        None,
        "星期六",
        None,
        None,
        "星期日",
        None,
        None,
        "合計",
    ]
    sheet.append(header_row)
    sheet.merge_cells(f"B{row_index}:D{row_index}")
    sheet.merge_cells(f"E{row_index}:G{row_index}")
    sheet.merge_cells(f"H{row_index}:J{row_index}")
    sheet.merge_cells(f"K{row_index}:M{row_index}")
    sheet.merge_cells(f"N{row_index}:P{row_index}")
    sheet.merge_cells(f"Q{row_index}:S{row_index}")
    sheet.merge_cells(f"T{row_index}:V{row_index}")

    for col_no in col_range:
        sheet[f"{col_no}{row_index}"].fill = light_gray
        sheet[f"{col_no}{row_index}"].border = border
        sheet[f"{col_no}{row_index}"].alignment = align_center
        sheet[f"{col_no}{row_index}"].font = bold

    row_index += 1
    header_row = [
        None,
        "" if table_widget.item(0, 1) is None else table_widget.item(0, 1).text(),
        None,
        None,
        "" if table_widget.item(0, 4) is None else table_widget.item(0, 4).text(),
        None,
        None,
        "" if table_widget.item(0, 7) is None else table_widget.item(0, 7).text(),
        None,
        None,
        "" if table_widget.item(0, 10) is None else table_widget.item(0, 10).text(),
        None,
        None,
        "" if table_widget.item(0, 13) is None else table_widget.item(0, 13).text(),
        None,
        None,
        "" if table_widget.item(0, 16) is None else table_widget.item(0, 16).text(),
        None,
        None,
        "" if table_widget.item(0, 19) is None else table_widget.item(0, 19).text(),
        None,
        None,
    ]
    sheet.append(header_row)
    sheet.merge_cells(f"B{row_index}:D{row_index}")
    sheet.merge_cells(f"E{row_index}:G{row_index}")
    sheet.merge_cells(f"H{row_index}:J{row_index}")
    sheet.merge_cells(f"K{row_index}:M{row_index}")
    sheet.merge_cells(f"N{row_index}:P{row_index}")
    sheet.merge_cells(f"Q{row_index}:S{row_index}")
    sheet.merge_cells(f"T{row_index}:V{row_index}")

    for col_no in col_range:
        sheet[f"{col_no}{row_index}"].border = border
        sheet[f"{col_no}{row_index}"].alignment = align_center
        sheet[f"{col_no}{row_index}"].font = bold

    row_index += 1
    header_row = [
        None,
        "早",
        "午",
        "晚",
        "早",
        "午",
        "晚",
        "早",
        "午",
        "晚",
        "早",
        "午",
        "晚",
        "早",
        "午",
        "晚",
        "早",
        "午",
        "晚",
        "早",
        "午",
        "晚",
        None,
    ]
    sheet.append(header_row)

    for col_no in col_range:
        sheet[f"{col_no}{row_index}"].fill = light_gray
        sheet[f"{col_no}{row_index}"].border = border
        sheet[f"{col_no}{row_index}"].alignment = align_center
        sheet[f"{col_no}{row_index}"].font = bold

    sheet.merge_cells("A2:A4")
    sheet.merge_cells("W2:W4")

    for row_no in range(2, table_widget.rowCount()):
        row = []
        for col_no in range(table_widget.columnCount()):
            row.append(
                ""
                if table_widget.item(row_no, col_no) is None
                else table_widget.item(row_no, col_no).text(),
            )
        sheet.append(row)

        row_index += 1
        for col in col_range:
            cell = sheet[f"{col}{row_index}"]
            cell.border = border
            cell.alignment = align_center
            if row_index % 2 == 0:
                cell.fill = alternate_row_color


def _export_week_doctor(workbook, title, clinic_name, year, month, table_widget):
    sheet = workbook.create_sheet(title)

    basic_col_range = [
        "A",
        "B",
        "C",
        "D",
        "E",
        "F",
        "G",
        "H",
        "I",
        "J",
        "K",
        "L",
        "M",
        "N",
        "O",
        "P",
    ]
    for col_no in basic_col_range:
        sheet.column_dimensions[col_no].width = 12

    row_index = 1
    sheet.row_dimensions[row_index].height = 30
    sheet.row_dimensions[row_index].alignment = align_left
    sheet.row_dimensions[row_index].font = bold
    sheet.append([f"{clinic_name} {year}年{month}月份 {title}"])
    sheet.merge_cells(f"A{row_index}:H{row_index}")

    row_index += 1
    col_range = []
    color_header = []
    for col_no in range(table_widget.columnCount()):
        col_range.append(basic_col_range[col_no])
        color_header.append(f"{basic_col_range[col_no]}{row_index}")

    header_row = []
    for i in range(table_widget.columnCount()):
        header_row.append(
            ""
            if table_widget.horizontalHeaderItem(i) is None
            else table_widget.horizontalHeaderItem(i).text()
        )
    sheet.append(header_row)

    for col_no in color_header:
        sheet[col_no].fill = light_gray
        sheet[col_no].border = border
        sheet[col_no].alignment = align_center
        sheet[col_no].font = bold

    for row_no in range(table_widget.rowCount()):
        row = []
        for col_no in range(table_widget.columnCount()):
            row.append(
                ""
                if table_widget.item(row_no, col_no) is None
                else table_widget.item(row_no, col_no).text(),
            )
        sheet.append(row)

        row_index += 1
        for col in col_range:
            cell = sheet[f"{col}{row_index}"]
            cell.border = border
            cell.alignment = align_center
            if row_index % 2 == 0:
                cell.fill = alternate_row_color


def export_medical_record_diagnosis_excel(
    database, excel_file_name, table_widget, title=None
):
    wb = Workbook()
    ws = wb.active
    ws.title = "sheet1"

    header_row = [
        "序",
        "就醫日期",
        "病歷號",
        "年齡",
        "性別",
        "主訴",
        "舌診",
        "脈象",
        "辨證",
        "治則",
        "病名",
    ]
    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 25
    ws.column_dimensions["C"].width = 10
    ws.column_dimensions["D"].width = 12
    ws.column_dimensions["E"].width = 5
    ws.column_dimensions["F"].width = 50
    ws.column_dimensions["G"].width = 20
    ws.column_dimensions["H"].width = 20
    ws.column_dimensions["I"].width = 20
    ws.column_dimensions["J"].width = 20
    ws.column_dimensions["K"].width = 40

    # if title is not None:
    #     ws.append([title])

    ws.append(header_row)

    sequence = 0
    for row_no in range(table_widget.rowCount()):
        sequence += 1

        row = []
        case_key = table_widget.item(row_no, 0).text()
        if case_key in [None, ""]:
            continue

        sql = f"""
            SELECT
                CaseDate, cases.PatientKey, cases.Name, patient.Gender, patient.Birthday,
                Symptom, Tongue, Pulse, Distincts, Cure, DiseaseCode1,
                DiseaseName1, DiseaseName2, DiseaseName3
            FROM cases
                LEFT JOIN patient ON patient.PatientKey = cases.PatientKey
            WHERE
                CaseKey = {case_key}
        """
        case_rows = database.select_record(sql)
        if len(case_rows) <= 0:
            continue

        case_row = case_rows[0]

        age_year, age_month = date_utils.get_age(
            case_row["Birthday"], case_row["CaseDate"]
        )
        if age_year is None:
            age = ""
        else:
            age = f"{age_year}歲{age_month}個月"

        row.append(sequence)
        row.append(case_row["CaseDate"])
        row.append(case_row["PatientKey"])
        row.append(age)
        row.append(string_utils.xstr(case_row["Gender"]))
        row.append(string_utils.get_str(case_row["Symptom"], "utf8").replace("\n", ""))
        row.append(string_utils.get_str(case_row["Tongue"], "utf8").replace("\n", ""))
        row.append(string_utils.get_str(case_row["Pulse"], "utf8").replace("\n", ""))
        row.append(
            string_utils.get_str(case_row["Distincts"], "utf8").replace("\n", "")
        )
        row.append(string_utils.get_str(case_row["Cure"], "utf8").replace("\n", ""))
        row.append(
            string_utils.xstr(case_row["DiseaseName1"])
            + "\n"
            + string_utils.xstr(case_row["DiseaseName2"])
            + "\n"
            + string_utils.xstr(case_row["DiseaseName3"])
        )

        try:
            ws.append(row)
            for medicine_set in range(1, 11):
                sql = f"""
                    SELECT PrescriptKey FROM prescript
                    WHERE
                        CaseKey = {case_key} AND
                        MedicineSet = {medicine_set}
                    LIMIT 1
                """
                row_count = database.select_record(sql)
                if len(row_count) <= 0:
                    continue

                if medicine_set == 1:
                    medicine_set_name = "健保處方"
                else:
                    medicine_set_name = f"自費處方{medicine_set - 1}"

                ws.append([medicine_set_name, ""])
                _write_treatment(database, ws, case_key, medicine_set)
                _write_medicine(database, ws, case_key, medicine_set)
        except Exception:
            pass

        ws.append([])

    for row_no, row in enumerate(ws.iter_rows()):
        symptom = row[5].value
        if symptom not in ["", None, "主訴"]:
            ws.row_dimensions[row_no + 1].height = 120

        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    wb.save(excel_file_name)
    try:
        # subprocess.Popen([excel_file_name], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


# 匯出醫療費用明細
def export_certificate_payment_to_excel(**kwargs):
    excel_file_name = kwargs["excel_file_name"]
    table_widget_certificate_items = kwargs["table_widget_certificate_items"]

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "醫療費用證明書"

    sheet.merge_cells("A1:M1")
    cell_telephone = sheet.cell(row=1, column=1)
    cell_telephone.value = f"{kwargs['clinic_name']} 醫療費用明細"
    cell_telephone.font = Font(size=20)
    cell_telephone.alignment = align_center
    sheet.row_dimensions[1].height = 50

    cell_telephone = sheet.cell(row=2, column=4)
    cell_telephone.value = f"電話: {kwargs['clinic_telephone']}"
    cell_telephone.font = Font(size=14)
    sheet.row_dimensions[2].height = 20

    cell_address = sheet.cell(row=3, column=4)
    cell_address.value = f"地址: {kwargs['clinic_address']}"
    cell_address.font = Font(size=14)
    sheet.row_dimensions[3].height = 20

    cell = sheet.cell(row=5, column=1)
    cell.value = f"收費期間: 從{kwargs['start_date']} 至 {kwargs['end_date']}"
    cell.font = Font(size=12)

    cell = sheet.cell(row=5, column=9)
    cell.value = f"列印日期: {date_utils.date_to_str()}"
    cell.font = Font(size=12)

    cell = sheet.cell(row=6, column=1)
    cell.value = f"病歷號碼: {kwargs['patient_key']}"
    cell.font = Font(size=12)

    cell = sheet.cell(row=6, column=5)
    cell.value = f"病患姓名: {kwargs['name']}"
    cell.font = Font(size=12)

    cell = sheet.cell(row=6, column=9)
    cell.value = f"身份證號: {kwargs['patient_id']}"
    cell.font = Font(size=12)

    cell = sheet.cell(row=7, column=1)
    cell.value = f"出生日期: {kwargs['birthday']}"
    cell.font = Font(size=12)

    cell = sheet.cell(row=7, column=5)
    cell.value = f"聯絡電話: {kwargs['telephone']}"
    cell.font = Font(size=12)

    cell = sheet.cell(row=8, column=1)
    cell.value = f"聯絡地址: {kwargs['address']}"
    cell.font = Font(size=12)

    sheet.append([])
    header_row = [
        "門診日期",
        "保險",
        "掛號費",
        "門診負擔",
        "藥品負擔",
        "小計",
        "健保申報",
        "自費藥費",
        "處置費",
        "其他費用",
        "折扣",
        "小計",
        "合計",
    ]
    sheet.append(header_row)

    sheet.column_dimensions["A"].width = 9
    sheet.column_dimensions["B"].width = 4
    sheet.column_dimensions["C"].width = 6
    sheet.column_dimensions["D"].width = 8
    sheet.column_dimensions["E"].width = 8
    sheet.column_dimensions["F"].width = 5
    sheet.column_dimensions["G"].width = 8
    sheet.column_dimensions["H"].width = 8
    sheet.column_dimensions["I"].width = 8
    sheet.column_dimensions["J"].width = 8
    sheet.column_dimensions["K"].width = 5
    sheet.column_dimensions["L"].width = 8
    sheet.column_dimensions["M"].width = 8

    for col_no in range(1, 14):  # head row
        cell = sheet.cell(row=10, column=col_no)
        cell.font = Font(size=9)
        cell.border = border

    start_no = 11
    last_row_no = start_no
    for row_no in range(table_widget_certificate_items.rowCount()):
        for col_no in range(2, 15):  # col 0, 1 = certificate_item_key, case_key
            value = table_widget_certificate_items.item(row_no, col_no).text()

            cell = sheet.cell(row=last_row_no, column=col_no - 1)
            if col_no >= 4:
                cell.value = number_utils.get_integer(value)
                cell.alignment = align_right
            else:
                cell.value = value

            cell.border = border
            cell.font = Font(size=9)

        cell = sheet.cell(row=last_row_no, column=6)
        cell.value = f"=SUM(C{last_row_no}:E{last_row_no})"

        cell = sheet.cell(row=last_row_no, column=12)
        cell.value = f"=SUM(H{last_row_no}:J{last_row_no})-K{last_row_no}"

        cell = sheet.cell(row=last_row_no, column=13)
        cell.value = f"=F{last_row_no}+G{last_row_no}+L{last_row_no}"

        last_row_no += 1

    total_fields = ["C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M"]
    col_no = 3
    current_row = last_row_no - 1
    for field in total_fields:
        cell = sheet.cell(row=current_row, column=col_no)
        cell.value = f"=SUM({field}{start_no}:{field}{current_row - 1})"
        col_no += 1

    sheet.merge_cells(f"A{current_row}:B{current_row}")
    cell = sheet.cell(row=current_row, column=1)
    cell.value = "合計"
    cell.alignment = align_center

    current_row += 1
    sheet.merge_cells(f"A{current_row}:M{current_row + 5}")
    cell = sheet.cell(row=current_row, column=1)
    cell.value = "本收據可為保險之憑證\n請妥善保存, 遺失恕不補發"
    cell.alignment = Alignment(horizontal="center", vertical="center", wrapText=True)
    cell.font = Font(size=14)

    current_row += 8
    cell = sheet.cell(row=current_row, column=2)
    cell.value = "大小章用印"
    cell.font = Font(size=12)

    sheet.merge_cells(f"E{current_row}:F{current_row + 4}")
    for i in range(5):
        cell = sheet[f"E{current_row + i}"]
        cell.border = border
        cell = sheet[f"F{current_row + i}"]
        cell.border = border

    current_row += 2
    sheet.merge_cells(f"H{current_row}:H{current_row + 2}")
    for i in range(3):
        cell = sheet[f"H{current_row + i}"]
        cell.border = border

    workbook.save(excel_file_name)

    try:
        # subprocess.Popen([excel_file_name], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


def _write_treatment(database, ws, case_key, medicine_set):
    sql = f"""
        SELECT * FROM prescript
        WHERE
            CaseKey = {case_key} AND
            MedicineSet = {medicine_set} AND
            MedicineType IN("穴道", "處置")
        ORDER BY MedicineSet, PrescriptKey
    """
    rows = database.select_record(sql)
    if len(rows) <= 0:
        return

    if string_utils.xstr(rows[0]["MedicineType"]) == "穴道":
        treatment = "針灸治療"
    else:
        treatment = "傷科治療"

    treatment_row = [
        "",
        treatment,
    ]
    ws.append(treatment_row)

    for row in rows:
        prescript_row = [
            "",
            string_utils.xstr(row["MedicineName"]),
        ]
        try:
            ws.append(prescript_row)
        except Exception:
            pass


def _write_medicine(database, ws, case_key, medicine_set):
    sql = f"""
        SELECT * FROM prescript
        WHERE
            CaseKey = {case_key} AND
            MedicineSet = {medicine_set} AND
            MedicineType NOT IN("穴道", "處置")
        ORDER BY MedicineSet, PrescriptKey
    """
    rows = database.select_record(sql)

    for row in rows:
        prescript_row = [
            "",
            string_utils.xstr(row["MedicineName"]),
            f"{row['Dosage']:.1f}{string_utils.xstr(row['Unit'])}",
        ]

        try:
            ws.append(prescript_row)
        except Exception:
            pass

    dosage_rows = case_utils.get_dosage_row(database, case_key, medicine_set)
    if len(dosage_rows) <= 0:
        return

    dosage_row = dosage_rows[0]

    packages = number_utils.get_integer(dosage_row["Packages"])
    days = number_utils.get_integer(dosage_row["Days"])
    if packages == 0 and days == 0:
        return

    instruction = dosage_row["Instruction"]
    dosage_line = f"1日{packages}包, 共{days}日份, {instruction}服用"

    try:
        ws.append(["", dosage_line])
    except Exception:
        pass


def export_list_to_excel(
    excel_file_name, header_list, list_name, numeric_cell=None, title=None
):
    if numeric_cell is None:
        numeric_cell = []

    wb = Workbook()
    ws = wb.active
    ws.title = "sheet1"

    header_row = []
    for col_no in range(len(header_list)):
        header_row.append(header_list[col_no])

    if title is not None:
        ws.append([title])

    ws.append(header_row)

    for row_no in range(len(list_name)):
        row = []
        for col_no in range(len(list_name[row_no])):
            item = list_name[row_no][col_no]
            if numeric_cell is not None and col_no in numeric_cell:
                item = item.replace(",", "")
                item = number_utils.get_float(item)

            row.append(item)

        ws.append(row)

    wb.save(excel_file_name)
    try:
        # subprocess.Popen([excel_file_name], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


# 匯出日報表 From medical_record_list 2019.07.01 板橋新生堂
def export_infectious_list(system_settings, excel_file_name, tableWidget_infectious):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "清冠一號藥品補助清冊"

    # sheet.merge_cells('A1:N1')
    sheet.merge_cells("A1:M1")
    cell = sheet.cell(row=1, column=1)
    cell.value = "醫療機構「公費清冠一號藥品費用」申請補助清冊"
    cell.alignment = align_center

    sheet.merge_cells("A2:A3")
    cell = sheet.cell(row=2, column=1)
    cell.value = "序號"
    cell.alignment = align_center

    sheet.merge_cells("B2:E2")
    cell = sheet.cell(row=2, column=2)
    cell.value = "院所資訊"
    cell.alignment = align_center

    sheet.merge_cells("F2:G2")
    cell = sheet.cell(row=2, column=6)
    cell.value = "個案資訊"
    cell.alignment = align_center

    sheet.merge_cells("H2:K2")
    cell = sheet.cell(row=2, column=8)
    cell.value = "藥品資訊"
    cell.alignment = align_center

    # sheet.merge_cells('L2:N2')
    sheet.merge_cells("L2:M2")
    cell = sheet.cell(row=2, column=12)
    cell.value = "院所查檢欄位"
    cell.alignment = align_center

    # sheet.merge_cells('B2:D2')
    # sheet.merge_cells('E2:F2')
    # sheet.merge_cells('G2:I2')
    # sheet.merge_cells('J2:L2')

    # header_row = [
    #     None, '醫事機構代碼', '醫療機構名稱', '申報費用年月', '申請費用', '姓名', '出生日期',
    #     '藥品品名', '用藥起始日期', '用藥天數', '開立總克數', 'COVID-19確診個案', '個案簽署同意書', '個案收置處所'
    # ]
    header_row = [
        None,
        "醫事機構代碼",
        "醫療機構名稱",
        "申報費用年月",
        "申請費用",
        "姓名",
        "出生日期",
        "藥品品名",
        "用藥起始日期",
        "用藥天數",
        "開立總克數",
        "COVID-19檢驗結果陽性",
        "同意書回傳院所",
    ]
    sheet.append(header_row)

    sheet.column_dimensions["A"].width = 8
    sheet.column_dimensions["B"].width = 16
    sheet.column_dimensions["C"].width = 30
    sheet.column_dimensions["D"].width = 15
    sheet.column_dimensions["E"].width = 12
    sheet.column_dimensions["F"].width = 15
    sheet.column_dimensions["G"].width = 15
    sheet.column_dimensions["H"].width = 30
    sheet.column_dimensions["I"].width = 15
    sheet.column_dimensions["J"].width = 12
    sheet.column_dimensions["K"].width = 13
    sheet.column_dimensions["L"].width = 22
    sheet.column_dimensions["M"].width = 20
    # sheet.column_dimensions['N'].width = 16

    for row_no in range(tableWidget_infectious.rowCount()):
        row = [row_no + 1]
        # for col_no in range(1, 14):
        for col_no in range(1, 13):
            value = tableWidget_infectious.item(row_no, col_no).text()
            if col_no in [4, 9]:
                value = number_utils.get_float(value)

            row.append(value)

        sheet.append(row)
        # for col_no in [1, 2, 3, 4, 5, 6, 7, 9, 10, 12, 13, 14]:
        for col_no in [1, 2, 3, 4, 5, 6, 7, 9, 10, 12, 13]:
            cell = sheet.cell(row=row_no + 4, column=col_no)
            cell.alignment = align_center

        for col_no in [11]:
            cell = sheet.cell(row=row_no + 4, column=col_no)
            cell.alignment = align_right

    last_row_no = tableWidget_infectious.rowCount() + 4

    sheet.merge_cells(f"A{last_row_no}:A{last_row_no}")
    cell = sheet.cell(row=last_row_no, column=2)
    cell.value = "填表人:"
    cell = sheet.cell(row=last_row_no, column=3)
    cell.value = system_utils.get_user_name(system_settings)

    sheet.merge_cells(f"A{last_row_no + 1}:A{last_row_no + 1}")
    cell = sheet.cell(row=last_row_no + 1, column=2)
    cell.value = "聯絡電話:"
    cell = sheet.cell(row=last_row_no + 1, column=3)
    cell.value = system_settings.field("院所電話")

    sheet.merge_cells(f"K{last_row_no}:K{last_row_no}")
    cell = sheet.cell(row=last_row_no, column=11)
    cell.value = "填表日期:"

    cell = sheet.cell(row=last_row_no, column=12)
    cell.value = date_utils.date_to_str()

    for row_no in range(1, sheet.max_row + 1):
        for col_no in range(1, sheet.max_column + 1):
            cell = sheet.cell(row=row_no, column=col_no)
            cell.border = border

    # start_no = 4
    # end_no = sheet.max_row-1
    # drug_list = [
    #     '“順天堂”RespireAid臺灣清冠一號濃縮顆粒',
    #     '“莊松榮”臺灣清冠一號濃縮顆粒',
    #     '康福顆粒(臺灣清冠一號)',
    #     '“勸奉堂”臺灣清冠一號濃縮顆粒',
    #     '“勝昌”臺灣清冠一號濃縮顆粒',
    #     '“華佗”臺灣清冠一號濃縮顆粒',
    #     '“漢聖”臺灣清冠一號濃縮顆粒',
    #     '“天一”臺灣清冠一號濃縮顆粒',
    #     '“科達”臺灣清冠一號濃縮顆粒',
    #     '“富田”臺灣清冠一號濃縮顆粒',
    # ]
    # _insert_combo_box_to_cell(sheet, drug_list, 'H', start_no, end_no)
    # _insert_combo_box_to_cell(sheet, ['☑', '☐'], 'L', start_no, end_no)
    # _insert_combo_box_to_cell(sheet, ['☑', '☐'], 'M', start_no, end_no)
    ## _insert_combo_box_to_cell(sheet, ['醫院', '集檢所', '防疫旅館', '居家'], 'N', start_no, end_no)

    workbook.save(excel_file_name)

    try:
        # subprocess.Popen([excel_file_name], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


def _insert_combo_box_to_cell(sheet, item_list, cell_alpha, start_no, end_no):
    list_str = ",".join(item_list)
    data_val = DataValidation(type="list", formula1=f'"{list_str}"', allow_blank=True)
    data_val.add(f"{cell_alpha}{start_no}:{cell_alpha}{end_no}")
    sheet.add_data_validation(data_val)


# 匯出交帳明細一覽表 2022.09.06 星光
def export_income_list(system_settings, excel_file_name, tableWidget_income, columns):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "交帳明細一覽表"

    clinic_name = system_settings.field("院所名稱")
    sheet.append([f"{clinic_name} 交帳明細一覽表"])
    sheet.merge_cells("A1:X1")
    row_property = sheet.row_dimensions[1]
    row_property.height = 30
    row_property.alignment = align_center
    row_property.font = bold

    header_row = [
        "序號",
        "門診日期",
        "班別",
        "病歷號",
        "姓名",
        "保險",
        "負擔類別",
        "就醫類別",
        "優待類別",
        "卡序",
        "藥日",
        "掛號費",
        "門診負擔",
        "藥品負擔",
        "欠卡費",
        "還卡費",
        "還款",
        "自費金額",
        "欠款",
        "掛號欠款",
        "民俗調理",
        "實收現金",
        "掛號人員",
        "批價人員",
    ]
    sheet.append(header_row)
    sheet.column_dimensions["A"].width = 6
    sheet.column_dimensions["B"].width = 14
    sheet.column_dimensions["C"].width = 7
    sheet.column_dimensions["D"].width = 10
    sheet.column_dimensions["E"].width = 12
    sheet.column_dimensions["F"].width = 7
    sheet.column_dimensions["G"].width = 20
    sheet.column_dimensions["H"].width = 20
    sheet.column_dimensions["I"].width = 20
    sheet.column_dimensions["J"].width = 12
    sheet.column_dimensions["K"].width = 7

    sheet.column_dimensions["L"].width = 10
    sheet.column_dimensions["M"].width = 10
    sheet.column_dimensions["N"].width = 10
    sheet.column_dimensions["O"].width = 10
    sheet.column_dimensions["P"].width = 10
    sheet.column_dimensions["Q"].width = 10
    sheet.column_dimensions["R"].width = 10
    sheet.column_dimensions["S"].width = 10
    sheet.column_dimensions["T"].width = 10
    sheet.column_dimensions["U"].width = 10
    sheet.column_dimensions["V"].width = 10
    sheet.column_dimensions["W"].width = 12
    sheet.column_dimensions["X"].width = 12

    for row_no in range(tableWidget_income.rowCount()):
        case_key_item = tableWidget_income.item(row_no, columns["case_key"])
        if case_key_item is None:
            continue

        if tableWidget_income.item(row_no, columns["name"]).text() == "合計":
            continue

        row = [row_no + 1]
        cell_no = row_no + 3
        for col_no in range(2, columns["cashier"] + 1):
            item = tableWidget_income.item(row_no, col_no)
            if item is None:
                value = None
            else:
                value = item.text()

            if col_no in [4, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22]:
                value = number_utils.get_integer(value)

            if col_no == columns["receipt_fee"]:
                value = f"=SUM(L{cell_no}:U{cell_no})"

            row.append(value)

        sheet.append(row)

    row = [
        None,
        None,
        None,
        None,
        "合計",
        None,
        None,
        None,
        None,
        None,
        None,
        f"=SUM(L3:L{cell_no})",
        f"=SUM(M3:M{cell_no})",
        f"=SUM(N3:N{cell_no})",
        f"=SUM(O3:O{cell_no})",
        f"=SUM(P3:P{cell_no})",
        f"=SUM(Q3:Q{cell_no})",
        f"=SUM(R3:R{cell_no})",
        f"=SUM(S3:S{cell_no})",
        f"=SUM(T3:T{cell_no})",
        f"=SUM(U3:U{cell_no})",
        f"=SUM(V3:V{cell_no})",
    ]
    sheet.append(row)
    sheet.freeze_panes = sheet["L2"]

    workbook.save(excel_file_name)

    try:
        # subprocess.Popen([excel_file_name], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


# 匯出診斷證明書 2023-09-09 陳立德的客戶
def export_certificate_diagnosis(
    database, system_settings, excel_file_name, certificate_key
):
    sql = f"""
        SELECT
            certificate.*,
            patient.Birthday, patient.Gender, patient.ID, patient.Telephone, patient.Cellphone,
            patient.Address
        FROM certificate
            LEFT JOIN patient ON patient.PatientKey = certificate.PatientKey
        WHERE
            CertificateKey = {certificate_key}
    """
    rows = database.select_record(sql)
    if len(rows) <= 0:
        return

    row = rows[0]

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "診斷證明書"

    sheet.merge_cells("A1:F1")
    cell_title = sheet.cell(row=1, column=1)
    cell_title.value = (
        f"{system_settings.field('院所名稱')} 診斷證明書\nCERTIFICATE OF DIAGNOSIS"
    )
    cell_title.font = Font(size=18)
    cell_title.alignment = align_center
    sheet.row_dimensions[1].height = 60

    sheet.row_dimensions[2].height = 40
    sheet.column_dimensions["A"].width = 13
    cell_name_header = sheet.cell(row=2, column=1)
    cell_name_header.value = "姓名\nName"
    cell_name_header.font = Font(size=12)
    cell_name_header.alignment = align_center
    cell_name_header.border = border
    # sheet.merge_cells('A2:B2')

    sheet.column_dimensions["B"].width = 13
    cell_name = sheet.cell(row=2, column=2)
    cell_name.value = string_utils.xstr(row["Name"])
    cell_name.font = Font(size=12)
    cell_name.alignment = align_center
    cell_name.border = border

    sheet.column_dimensions["C"].width = 13
    cell_gender_header = sheet.cell(row=2, column=3)
    cell_gender_header.value = "性別\nGender"
    cell_gender_header.font = Font(size=12)
    cell_gender_header.alignment = align_center
    cell_gender_header.border = border

    sheet.column_dimensions["D"].width = 14
    cell_gender = sheet.cell(row=2, column=4)
    cell_gender.value = string_utils.xstr(row["Gender"])
    cell_gender.font = Font(size=12)
    cell_gender.alignment = align_center
    cell_gender.border = border

    sheet.column_dimensions["E"].width = 16
    cell_birth_header = sheet.cell(row=2, column=5)
    cell_birth_header.value = "出生日期\nDate of Birth"
    cell_birth_header.font = Font(size=12)
    cell_birth_header.alignment = align_center
    cell_birth_header.border = border

    birthday = string_utils.xstr(row["Birthday"])
    tw_birthday = date_utils.west_date_to_nhi_date(row["Birthday"], "-")

    sheet.column_dimensions["F"].width = 17
    cell_birth = sheet.cell(row=2, column=6)
    cell_birth.value = f"{birthday}\n(民國{tw_birthday})"
    cell_birth.font = Font(size=12)
    cell_birth.alignment = align_center
    cell_birth.border = border

    sheet.row_dimensions[3].height = 40
    chart_no_header = sheet.cell(row=3, column=1)
    chart_no_header.value = "病歷號碼\nChart No."
    chart_no_header.font = Font(size=12)
    chart_no_header.alignment = align_center
    chart_no_header.border = border

    chart_no = sheet.cell(row=3, column=2)
    chart_no.value = f"{row['PatientKey']:0>6}"
    chart_no.font = Font(size=12)
    chart_no.alignment = align_center
    chart_no.border = border

    id_no_header = sheet.cell(row=3, column=3)
    id_no_header.value = "身份證號\nID No."
    id_no_header.font = Font(size=12)
    id_no_header.alignment = align_center
    id_no_header.border = border

    id_no = sheet.cell(row=3, column=4)
    id_no.value = f"{string_utils.xstr(row['ID'])}"
    id_no.font = Font(size=12)
    id_no.alignment = align_center
    id_no.border = border

    phone_header = sheet.cell(row=3, column=5)
    phone_header.value = "電話\nPhone No."
    phone_header.font = Font(size=12)
    phone_header.alignment = align_center
    phone_header.border = border

    telephone = string_utils.xstr(row["Telephone"])
    if telephone == "":
        telephone = string_utils.xstr(row["Cellphone"])

    phone = sheet.cell(row=3, column=6)
    phone.value = telephone
    phone.font = Font(size=12)
    phone.alignment = align_center
    phone.border = border

    sheet.row_dimensions[4].height = 40
    address_header = sheet.cell(row=4, column=1)
    address_header.value = "地址\nAddress"
    address_header.font = Font(size=12)
    address_header.alignment = align_center
    address_header.border = border

    address = sheet.cell(row=4, column=2)
    address.value = string_utils.xstr(row["Address"])
    address.font = Font(size=12)
    address.alignment = align_center
    address.border = border
    sheet.merge_cells("B4:F4")

    sheet.row_dimensions[5].height = 40
    class_header = sheet.cell(row=5, column=1)
    class_header.value = "科別\nSpeciality"
    class_header.font = Font(size=12)
    class_header.alignment = align_center
    class_header.border = border

    class_type = sheet.cell(row=5, column=2)
    class_type.value = "60 中醫科"
    class_type.font = Font(size=12)
    class_type.alignment = align_center
    class_type.border = border

    case_date_header = sheet.cell(row=5, column=3)
    case_date_header.value = "診療日期\nExam Date"
    case_date_header.font = Font(size=12)
    case_date_header.alignment = align_center
    case_date_header.border = border

    case_date = date_utils.date_to_zh_tw_date(string_utils.xstr(row["StartDate"]))
    if row["EndDate"] != row["StartDate"]:
        end_date = date_utils.date_to_zh_tw_date(string_utils.xstr(row["EndDate"]))
        case_date += f" 至 {end_date}"

    case_list = get_case_list(database, certificate_key, row)
    case_times = len(case_list)

    case_date_header = sheet.cell(row=5, column=4)
    case_date_header.value = f"{case_date}\n共{case_times}次"
    case_date_header.font = Font(size=12)
    case_date_header.alignment = align_center
    case_date_header.border = border
    sheet.merge_cells("D5:F5")

    sheet.row_dimensions[6].height = 20
    diagnosis_header = sheet.cell(row=6, column=1)
    diagnosis_header.value = "診斷 Diagnosis"
    diagnosis_header.font = Font(size=12)
    diagnosis_header.alignment = align_left
    diagnosis_header.border = border
    sheet.merge_cells("A6:F6")

    sheet.row_dimensions[7].height = 80
    diagnosis = sheet.cell(row=7, column=1)
    diagnosis.value = string_utils.xstr(row["Diagnosis"])
    diagnosis.font = Font(size=12)
    diagnosis.alignment = align_left
    diagnosis.border = border
    sheet.merge_cells("A7:F7")

    sheet.row_dimensions[8].height = 20
    comment_header = sheet.cell(row=8, column=1)
    comment_header.value = "醫囑 Doctor's Comment"
    comment_header.font = Font(size=12)
    comment_header.alignment = align_left
    comment_header.border = border
    sheet.merge_cells("A8:F8")

    sheet.row_dimensions[9].height = 80
    comment = sheet.cell(row=9, column=1)
    comment.value = string_utils.xstr(row["DoctorComment"])
    comment.font = Font(size=12)
    comment.alignment = align_left
    comment.border = border
    sheet.merge_cells("A9:F9")

    year = row["CertificateDate"].year
    month = row["CertificateDate"].month
    day = row["CertificateDate"].day
    certificate_date = f"{year} 年 {month} 月 {day} 日"

    physician = string_utils.xstr(row["Doctor"])
    physician_cert_no = personnel_utils.get_person_field_value(
        database, physician, "Certificate"
    )
    president = system_settings.field("負責醫師")
    license_no = system_settings.field("院所代號")
    clinic_telephone = system_settings.field("院所電話")
    clinic_address = system_settings.field("院所地址")

    sheet.row_dimensions[10].height = 250
    summary = sheet.cell(row=10, column=1)
    summary.value = f"""以上病人經本院(所)醫師診斷屬實特予證明\n
        主治醫師: {physician}
        醫師證書號碼: {physician_cert_no}

        院長: {president}
        開業執照號碼: {license_no}

        院所電話: {clinic_telephone}
        院所地址: {clinic_address}

        開立診斷證明書日期: {certificate_date}
    """
    summary.font = Font(size=12)
    summary.alignment = Alignment(horizontal="left", vertical="top")
    summary.border = border
    sheet.merge_cells("A10:F10")

    sheet.row_dimensions[11].height = 24
    note = sheet.cell(row=11, column=1)
    note.value = "本證明書經塗改或未加蓋本院印章值者無效\n本證明書訴訟無效"
    note.font = Font(size=8)
    note.alignment = align_left
    sheet.merge_cells("A11:F11")

    workbook.save(excel_file_name)

    try:
        # subprocess.Popen([excel_file_name], shell=True)
        open_file(excel_file_name)
    except Exception:
        pass


def get_case_list(database, certificate_key, row):
    sql = f"""
        SELECT CaseDate FROM certificate_items
        WHERE
            CertificateKey = {certificate_key}
        ORDER BY CaseDate
    """
    rows = database.select_record(sql)

    if len(rows) <= 0:
        rows = _get_rows_by_script(database, row)

    case_list = []
    for row in rows:
        case_list.append(string_utils.xstr(row["CaseDate"].date()))

    return case_list


def _get_rows_by_script(database, row):
    start_date = row["StartDate"]
    start_date = f"{start_date} 00:00:00"

    end_date = row["EndDate"]
    end_date = f"{end_date} 23:59:59"

    patient_key = row["PatientKey"]

    treat_type_dict = {
        "針傷科": nhi_utils.INS_TREAT,
        "針灸科": nhi_utils.ACUPUNCTURE_TREAT,
        "傷骨科": nhi_utils.MASSAGE_TREAT,
    }

    condition = ""
    ins_type = string_utils.xstr(row["InsType"])
    treat_type = string_utils.xstr(row["TreatType"])
    if treat_type == "":
        treat_type = "全部"

    if ins_type in ["健保", "自費"]:
        condition = f' AND InsType = "{ins_type}" '

    if treat_type == "內科":
        condition += ' AND TreatType = "內科" '
    elif treat_type != "全部":
        treat_type_list = tuple(treat_type_dict[treat_type])
        condition += f" AND TreatType IN {treat_type_list} "

    sql = f'''
        SELECT CaseDate FROM cases
        WHERE
            CaseDate BETWEEN "{start_date}" AND "{end_date}" AND
            PatientKey = {patient_key} AND
            TreatType != "自購"
            {condition}
        GROUP BY DATE(CaseDate)
        ORDER BY CaseDate
    '''
    rows = database.select_record(sql)

    return rows


# 匯出病歷至Word
def export_case_to_word(database, filename, case_key_list):
    try:
        pass
    except Exception:
        system_utils.pip3_install("python-docx")

    from docx import Document

    doc = Document()

    for case_key in case_key_list:
        export_case_to_word_by_case_key(database, doc, case_key)

    doc.save(filename)


def export_case_to_word_by_case_key(database, doc, case_key):
    sql = f"""
        SELECT cases.*, patient.Gender, patient.Birthday FROM cases
            LEFT JOIN patient ON patient.PatientKey = cases.PatientKey
        WHERE
            CaseKey = {case_key}
    """

    rows = database.select_record(sql)
    if not rows:
        return

    row = rows[0]
    card = string_utils.xstr(row["Card"])
    continuance = number_utils.get_integer(row["Continuance"])
    if continuance >= 1:
        card += f"-{continuance}"

    age, _ = date_utils.get_age(row["Birthday"], row["CaseDate"])
    doc.add_heading(
        f"""門診日{row["CaseDate"].date()}\t保險: {row["InsType"]} {row["Share"]} {card}\t醫師: {row["Doctor"]}""",
        level=1,
    )
    doc.add_heading(
        f"姓名: {row['Name']}\t性別: {row['Gender']}\t生日: {row['Birthday']}\t年齡: {age}歲",
        level=2,
    )

    disease_code1 = string_utils.xstr(row["DiseaseCode1"])
    disease_code2 = string_utils.xstr(row["DiseaseCode2"])
    disease_code3 = string_utils.xstr(row["DiseaseCode3"])
    if disease_code1 != "":
        doc.add_heading(f"""主診斷碼:{disease_code1} {row["DiseaseName1"]}""", level=3)
    if disease_code2 != "":
        doc.add_heading(f"""次斷碼1:{disease_code2} {row["DiseaseName2"]}""", level=3)
    if disease_code3 != "":
        doc.add_heading(f"""次斷碼2:{disease_code3} {row["DiseaseName3"]}""", level=3)

    symptom = string_utils.get_str(row["Symptom"], "utf-8")
    symptom = symptom.replace("\n", "")
    symptom = symptom.replace("\t", "")
    if symptom != "":
        doc.add_heading(f"主訴: {symptom}", level=3)

    tongue = string_utils.get_str(row["Tongue"], "utf-8")
    if tongue != "":
        doc.add_heading(f"舌診: {tongue}", level=3)

    pulse = string_utils.get_str(row["Pulse"], "utf-8")
    if pulse != "":
        doc.add_heading(f"脈象: {pulse}", level=3)

    distinct = string_utils.get_str(row["Distincts"], "utf-8")
    if distinct != "":
        doc.add_heading(f"辨證: {distinct}", level=3)

    cure = string_utils.get_str(row["Cure"], "utf-8")
    if cure != "":
        doc.add_heading(f"治則: {cure}", level=3)

    remark = string_utils.get_str(row["Remark"], "utf-8")
    if remark != "":
        doc.add_heading(f"備註: {remark}", level=3)

    treatment = string_utils.xstr(row["Treatment"])
    export_prescript_to_word_by_case_key(database, doc, case_key, treatment)
    doc.add_heading(
        "-----------------------------------------------------------------------------------"
    )
    # add_horizontal_line(doc)


def set_cell_background(cell, color_hex: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    """設定儲存格背景顏色（例如 'D9D9D9' 為淡灰色）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def export_prescript_to_word_by_case_key(database, doc, case_key, treatment):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    sql = f"""
        SELECT * FROM prescript
        WHERE
            CaseKey = {case_key}
        GROUP By MedicineSet
        ORDER BY MedicineSet
    """
    rows = database.select_record(sql)
    for row in rows:
        doc.add_paragraph()

        medicine_set = row["MedicineSet"]
        if medicine_set == 1:
            ins_type = "健保"
        else:
            ins_type = f"自費({medicine_set - 1})"

        total_width = 6  # 單位：Inches
        # 填入表頭
        if medicine_set == 1:
            table = doc.add_table(rows=1, cols=3)  # 建立 1 列 3 欄的表格
            # 設定每個 cell 的寬度（模擬百分比）
            headers = [ins_type + "處方", "劑量", "單位"]
            cell_widths = [0.6, 0.2, 0.2]  # 百分比表示（60%、20%、20%）
        else:
            table = doc.add_table(rows=1, cols=5)
            headers = [ins_type + "處方", "劑量", "單位", "單價", "金額"]
            cell_widths = [0.5, 0.1, 0.1, 0.15, 0.15]

        for i, width_ratio in enumerate(cell_widths):
            table.columns[i].width = Inches(total_width * width_ratio)

        table.style = "Table Grid"  # 設定表格樣式（讓格線可見）
        hdr_cells = table.rows[0].cells

        for i, text in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = text
            set_cell_background(cell, "D9D9D9")  # 淡灰色
            # 控制字型與對齊
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.bold = True

        sql = f"""
            SELECT * FROM prescript
            WHERE
                CaseKey = {case_key} AND
                MedicineSet = {medicine_set}
            ORDER BY PrescriptNo
        """
        prescript_rows = database.select_record(sql)
        data = []
        if treatment not in ["", None]:
            data.append([treatment, "", ""])

        for prescript_row in prescript_rows:
            try:
                dosage = f"{prescript_row['Dosage']:.1f}"
            except Exception:
                dosage = ""

            if medicine_set == 1:
                data.append(
                    [
                        string_utils.xstr(prescript_row["MedicineName"]),
                        dosage,
                        string_utils.xstr(prescript_row["Unit"]),
                    ]
                )
            else:
                data.append(
                    [
                        string_utils.xstr(prescript_row["MedicineName"]),
                        dosage,
                        string_utils.xstr(prescript_row["Unit"]),
                        string_utils.xstr(prescript_row["Price"]),
                        string_utils.xstr(prescript_row["Amount"]),
                    ]
                )

        for row_data in data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = value
                if (medicine_set == 1 and i == 1) or (
                    medicine_set >= 2 and i in [1, 3, 4]
                ):
                    for paragraph in row_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif i == 2:
                    for paragraph in row_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        packages = case_utils.get_packages(database, case_key, medicine_set)
        pres_days = case_utils.get_pres_days(database, case_key, medicine_set)
        instuction = case_utils.get_instruction(database, case_key, medicine_set)
        if packages > 0 or pres_days > 0:
            instuction_str = f"{packages}包{pres_days}天 服用方式: {instuction}"
            # 新增一列，合併三個儲存格
            instruction_row = table.add_row().cells
            if medicine_set == 1:
                merged_cell = (
                    instruction_row[0]
                    .merge(instruction_row[1])
                    .merge(instruction_row[2])
                )
            else:
                merged_cell = (
                    instruction_row[0]
                    .merge(instruction_row[1])
                    .merge(instruction_row[4])
                )

            merged_cell.text = instuction_str

            # 可選：設定文字樣式與置中
            for paragraph in merged_cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(12)


def add_horizontal_line(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # 取得段落屬性的 <w:pPr> 標籤
    p_pr = p._p.get_or_add_pPr()

    # 建立段落邊框設定 <w:pBdr>
    p_borders = OxmlElement("w:pBdr")

    # 建立底線 <w:bottom>
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")  # 線條樣式：single、double、dotted 等
    bottom.set(qn("w:sz"), "6")  # 線條粗細（1/8 pt 單位）→ 6 = 0.75pt
    bottom.set(qn("w:space"), "1")  # 線條與文字間距
    bottom.set(qn("w:color"), "000000")  # 黑色

    # 加入到底線元素
    p_borders.append(bottom)
    p_pr.append(p_borders)

    return p


def export_correction_reg_income_txt(
    database, system_settings, txt_filename, tableWidget_medical_record_list, column
):
    clinic_id = system_settings.field("院所代號")

    lines = []
    for row_no in range(tableWidget_medical_record_list.rowCount()):
        try:
            check_box_print_mark = tableWidget_medical_record_list.cellWidget(
                row_no, column["PrintMark"]
            )
            if not check_box_print_mark.isChecked():
                continue
        except Exception:
            continue

        case_key = tableWidget_medical_record_list.item(
            row_no, column["CaseKey"]
        ).text()
        sql = f"""
            SELECT
                cases.CaseDate, cases.DiseaseCode1, cases.RegistFee,
                cases.SDiagShareFee, cases.SDrugShareFee,
                patient.ID, patient.CardNo, patient.Birthday
            FROM cases
                LEFT JOIN patient ON patient.PatientKey = cases.PatientKey
            WHERE
                cases.CaseKey = {case_key}
        """
        case_row = database.select_record(sql)
        if not case_row:
            continue

        case_row = case_row[0]
        card_no = case_row["CardNo"]
        pid = case_row["ID"]
        birth_date = date_utils.west_date_to_nhi_date(case_row["Birthday"])
        case_date = date_utils.west_datetime_to_nhi_datetime(case_row["CaseDate"])
        disease_code = string_utils.xstr(case_row["DiseaseCode1"])[:5].ljust(5)
        regist_fee = string_utils.xstr(case_row["RegistFee"]).ljust(8)
        share_fee = number_utils.get_integer(
            case_row["SDiagShareFee"]
        ) + number_utils.get_integer(case_row["SDrugShareFee"])
        share_fee = string_utils.xstr(share_fee).ljust(8)
        hosp_fee1 = "0".ljust(7)
        hosp_fee2 = "0".ljust(7)
        drug_fee = "0".ljust(7)
        material_fee = "0".ljust(7)
        self_fee = "0".ljust(7)
        description = " ".ljust(100)
        lines.append(
            f"{card_no}{pid}{birth_date}{clinic_id}"
            f"{case_date}{disease_code}{regist_fee}{share_fee}"
            f"{hosp_fee1}{hosp_fee2}{drug_fee}{material_fee}{self_fee}"
            f"{description}\n"
        )

    with open(txt_filename, "w", encoding="big5") as f:
        for line in lines:
            f.write(line)
